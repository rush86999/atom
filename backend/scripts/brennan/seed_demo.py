"""
Brennan.ca demo seeder.

Writes brennan-specific entities into the knowledge graph (GraphNode/GraphEdge) and
generates branded Office templates (.docx/.xlsx) with placeholder tokens.

Usage:
    cd backend && venv/bin/python scripts/brennan/seed_demo.py

Idempotent: re-running updates existing entities/templates in place.
"""

import os
import sys
from datetime import datetime
from pathlib import Path

# Make the demo data importable and the backend package importable.
REPO_ROOT = Path(__file__).resolve().parents[3]
DEMO_DIR = REPO_ROOT / "demo" / "brennan"
sys.path.insert(0, str(DEMO_DIR))
sys.path.insert(0, str(REPO_ROOT / "backend"))

from seed_data import (  # noqa: E402
    WORKSPACE_ID,
    PRODUCTS,
    CUSTOMERS,
    VENDORS,
    LEADS,
    RELATIONSHIPS,
    WORD_PLACEHOLDERS,
    EXCEL_TEMPLATES,
)


def _ensure_graph_columns() -> None:
    """Ensure graph_nodes has the columns the ORM model expects.

    Older SQLite DBs created via Base.metadata.create_all may be missing the
    `embedding` column (added to the model after initial table creation). Without
    it, GraphRAGEngine.ingest_structured_data raises OperationalError. We add it
    idempotently here so the demo seeder is self-contained.
    """
    try:
        from sqlalchemy import inspect, text
        from core.database import engine

        insp = inspect(engine)
        if not insp.has_table("graph_nodes"):
            return
        existing_cols = {c["name"] for c in insp.get_columns("graph_nodes")}
        with engine.begin() as conn:
            if "embedding" not in existing_cols:
                conn.execute(text("ALTER TABLE graph_nodes ADD COLUMN embedding JSON"))
                print("  (added missing 'embedding' column to graph_nodes)")
    except Exception as e:
        print(f"  (could not verify graph_nodes schema: {e})")


def seed_knowledge_graph() -> dict:
    """Seed products, contacts, leads, and relationships into the knowledge graph.

    Uses GraphRAGEngine.ingest_structured_data — the same path the agent's ingestion
    fetchers and the Outlook learner use — so seeded entities are queryable the same
    way as real ingested data.
    """
    _ensure_graph_columns()
    from core.graphrag_engine import GraphRAGEngine

    # Build entity dicts in the shape ingest_structured_data expects:
    # {name, type, description, properties}.
    entities = []
    for ent in PRODUCTS + CUSTOMERS + VENDORS + LEADS:
        props = dict(ent.get("properties", {}))
        entities.append(
            {
                "name": ent["name"],
                "type": ent["type"],
                "description": props.get("description", ""),
                "properties": props,
            }
        )

    engine = GraphRAGEngine()
    result = engine.ingest_structured_data(
        workspace_id=WORKSPACE_ID,
        entities=entities,
        relationships=list(RELATIONSHIPS),
    )
    return {
        "entities": result.get("entities", 0),
        "relationships": result.get("relationships", 0),
    }


def _generate_word_templates(template_dir: Path) -> int:
    """Create .docx templates containing literal placeholder tokens."""
    from docx import Document

    count = 0
    for fname, tokens in WORD_PLACEHOLDERS.items():
        out_path = template_dir / fname
        doc = Document()

        # Header
        doc.add_heading("BRENNAN.CA", level=0)
        doc.add_paragraph("Metal Fabrication Machinery & Parts")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')} (demo template)")

        doc.add_heading(fname.replace(".docx", ""), level=1)

        # Body: list each placeholder token on its own line so the agent can
        # target it with modify_word_document(action="replace", target=token).
        doc.add_paragraph("---- Fill the fields below ----")
        for token in tokens:
            doc.add_paragraph(token)

        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        count += 1
        print(f"  ✓ {fname} ({len(tokens)} placeholders)")

    return count


def _generate_excel_templates(template_dir: Path) -> int:
    """Create .xlsx templates with named sheets and seed rows."""
    from openpyxl import Workbook

    count = 0
    for fname, spec in EXCEL_TEMPLATES.items():
        out_path = template_dir / fname
        wb = Workbook()
        ws = wb.active
        ws.title = spec["sheet"]

        for cell_path, value, is_formula in spec["cells"]:
            # cell_path is like /SheetName/A1 — extract the coordinate.
            parts = cell_path.strip("/").split("/")
            coordinate = parts[-1]
            if is_formula and isinstance(value, str):
                ws[coordinate] = value
            else:
                ws[coordinate] = value

        out_path.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(out_path))
        count += 1
        print(f"  ✓ {fname} ({len(spec['cells'])} cells)")

    return count


def generate_templates() -> int:
    """Generate all Office templates into demo/brennan/templates/."""
    template_dir = DEMO_DIR / "templates"
    template_dir.mkdir(parents=True, exist_ok=True)
    print("Generating Office templates...")
    n = _generate_word_templates(template_dir) + _generate_excel_templates(template_dir)
    return n


def main():
    print("=" * 60)
    print("Brennan.ca Demo Seeder")
    print("=" * 60)

    # 1. Knowledge graph
    print("\n[1/2] Seeding knowledge graph (products, contacts, leads)...")
    try:
        counts = seed_knowledge_graph()
        print(
            f"  ✓ {counts['entities']} entities + {counts['relationships']} relationships "
            f"written to workspace '{WORKSPACE_ID}'"
        )
    except Exception as e:
        print(f"  ✗ Knowledge graph seeding failed: {e}")
        print("    (Ensure the backend DB is initialized — run the server once first.)")

    # 2. Templates
    print("\n[2/2] Generating Office templates...")
    try:
        n = generate_templates()
        print(f"  ✓ {n} templates written to {DEMO_DIR / 'templates'}")
    except Exception as e:
        print(f"  ✗ Template generation failed: {e}")

    print("\nDone. Next steps:")
    print(f"  • Verify with: venv/bin/python scripts/brennan/verify_memory.py")
    print(f"  • Templates are in: {DEMO_DIR / 'templates'}")
    print(f"  • Storyboards are in: {DEMO_DIR / 'storyboards'}")


if __name__ == "__main__":
    main()
