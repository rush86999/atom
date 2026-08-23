"""Coverage wave 61 — core/office_sync_service.py (93% → 98%+).

Closes the remaining gaps: broadcast_file_update outer exception path,
docx broadcast component mapping (docs/rich_editor), sync ingest missing-file
and exception paths, _read_file_bytes read-failure path, empty-file handling,
ingest status edge cases, and the event-loop leak in the sync ingest fallback
(temp loop must be closed after run_until_complete — was leaked, open selector
FD per sync-context office edit).
"""
import asyncio
import os
import tempfile
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock, Mock, patch

import openpyxl
import pytest

from core.office_sync_service import OfficeSyncService


@pytest.fixture
def svc(tmp_path, monkeypatch):
    monkeypatch.setenv("ATOM_OFFICE_DIR", str(tmp_path))
    db = Mock()
    return OfficeSyncService(db), db, tmp_path


def _make_xlsx(tmp_path, name="a.xlsx"):
    p = tmp_path / name
    wb = openpyxl.Workbook()
    wb.active.title = "Sheet1"
    wb.save(p)
    return p


class TestSyncCanvasToFile:
    def test_invalid_path_contained(self, svc):
        s, db, _ = svc
        result = s.sync_canvas_to_file("c1", "/etc/passwd", "u1", "cell", {})
        assert result["success"] is False
        assert "Access denied" in result["error"]

    def test_missing_file(self, svc, tmp_path):
        s, db, _ = svc
        result = s.sync_canvas_to_file(
            "c1", str(tmp_path / "missing.xlsx"), "u1", "cell", {})
        assert result["success"] is False
        assert "not found" in result["error"]

    def test_cell_path_required(self, svc, tmp_path):
        s, db, _ = svc
        p = _make_xlsx(tmp_path)
        result = s.sync_canvas_to_file("c1", str(p), "u1", "cell", {})
        assert result["success"] is False
        assert "cell_path required" in result["error"]

    def test_write_failure_propagated(self, svc, tmp_path):
        s, db, _ = svc
        p = _make_xlsx(tmp_path)
        with patch.object(s.office.excel, "write_cell",
                          return_value={"success": False, "error": "bad"}):
            result = s.sync_canvas_to_file(
                "c1", str(p), "u1", "cell", {"cell_path": "/S/A1", "value": 1})
        assert result["success"] is False
        assert result["error"] == "bad"

    def test_docx_content_rewrite(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "d.docx"
        import docx
        doc = docx.Document()
        doc.add_paragraph("old")
        doc.save(p)
        with patch.object(s, "broadcast_file_update") as bcast:
            result = s.sync_canvas_to_file(
                "c1", str(p), "u1", "document", {"content": "l1\nl2"})
        assert result["success"] is True
        bcast.assert_called_once()
        assert [pa.text for pa in docx.Document(p).paragraphs] == ["l1", "l2"]

    def test_unsupported_edit_type(self, svc, tmp_path):
        s, db, _ = svc
        p = _make_xlsx(tmp_path)
        result = s.sync_canvas_to_file(
            "c1", str(p), "u1", "document", {"content": "x"})
        assert result["success"] is False
        assert "Unsupported sync edit type" in result["error"]

    def test_xlsx_with_formula_flag(self, svc, tmp_path):
        s, db, _ = svc
        p = _make_xlsx(tmp_path)
        with patch.object(s.office.excel, "write_cell",
                          return_value={"success": True}) as wc, \
             patch.object(s, "broadcast_file_update"):
            result = s.sync_canvas_to_file(
                "c1", str(p), "u1", "cell",
                {"cell_path": "/Sheet1/A1", "value": "=1+1", "is_formula": True})
        assert result["success"] is True
        wc.assert_called_once()
        kwargs = wc.call_args[1]
        assert kwargs["is_formula"] is True

    def test_explicit_none_content_docx_fails_cleanly(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "d.docx"
        import docx
        doc = docx.Document()
        doc.add_paragraph("x")
        doc.save(p)
        with patch.object(s, "broadcast_file_update"):
            result = s.sync_canvas_to_file(
                "c1", str(p), "u1", "document", {"content": None})
        assert result["success"] is False


class TestBroadcast:
    def test_invalid_path_returns(self, svc):
        s, db, _ = svc
        s.broadcast_file_update("c1", "/etc/passwd", "u1")

    def test_render_failure_returns(self, svc, tmp_path):
        s, db, _ = svc
        p = _make_xlsx(tmp_path)
        with patch.object(s.office.renderer, "render_to_html",
                          return_value={"success": False}):
            s.broadcast_file_update("c1", str(p), "u1")
        # Render failure degrades to html=None — the audit row is still
        # written (structured snapshot is independent of the render).
        audit = db.add.call_args[0][0]
        assert audit.details_json["html"] is None

    def test_render_exception_swallowed(self, svc, tmp_path):
        """Outer except in broadcast_file_update: render raising must not propagate."""
        s, db, _ = svc
        p = _make_xlsx(tmp_path)
        with patch.object(s.office.renderer, "render_to_html",
                          side_effect=RuntimeError("render exploded")):
            s.broadcast_file_update("c1", str(p), "u1")
        db.add.assert_not_called()
        db.commit.assert_not_called()

    def test_db_commit_exception_swallowed(self, svc, tmp_path):
        s, db, _ = svc
        p = _make_xlsx(tmp_path)
        db.commit.side_effect = RuntimeError("commit failed")
        with patch.object(s.office.renderer, "render_to_html",
                          return_value={"success": True, "html": "<b>x</b>"}), \
             patch.object(s, "_ingest_document_to_memory",
                          new=AsyncMock(return_value=True)), \
             patch("core.office_sync_service.ws_manager.broadcast",
                   new=AsyncMock()):
            s.broadcast_file_update("c1", str(p), "u1")
        db.add.assert_called_once()

    def test_docx_broadcast_component_mapping(self, svc, tmp_path):
        """docx ext must map to docs/rich_editor (else-branch of component map)."""
        s, db, _ = svc
        p = tmp_path / "d.docx"
        import docx
        doc = docx.Document()
        doc.add_paragraph("hello")
        doc.save(p)
        with patch.object(s.office.renderer, "render_to_html",
                          return_value={"success": True, "html": "<p>hello</p>"}), \
             patch.object(s, "_ingest_document_to_memory",
                          new=AsyncMock(return_value=True)), \
             patch.object(s, "_ingest_document_to_memory_sync", return_value=True), \
             patch("core.office_sync_service.ws_manager.broadcast",
                   new=AsyncMock()) as wsb:
            s.broadcast_file_update("c1", str(p), "u1")
        audit = db.add.call_args[0][0]
        assert audit.canvas_type == "docs"
        assert audit.details_json["component_type"] == "rich_editor"
        assert audit.details_json["title"] == "d.docx"
        wsb.assert_called()


class TestIngestSync:
    def test_missing_file_returns_false(self, svc, tmp_path):
        s, db, _ = svc
        assert s._ingest_document_to_memory_sync(str(tmp_path / "nope"), "u1") is False

    def test_exception_returns_false(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "f.txt"
        p.write_bytes(b"content")
        with patch("core.auto_document_ingestion.AutoDocumentIngestionService",
                   side_effect=RuntimeError("boom")):
            assert s._ingest_document_to_memory_sync(str(p), "u1") is False

    def test_failed_status_returns_false(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "f.txt"
        p.write_bytes(b"content")
        with patch("core.auto_document_ingestion.AutoDocumentIngestionService") as cls:
            ing = cls.return_value
            ing.process_file_bytes = AsyncMock(return_value={"status": "failed"})
            assert s._ingest_document_to_memory_sync(str(p), "u1") is False

    def test_success_closes_temp_event_loop(self, svc, tmp_path):
        """Regression: the temp event loop created for sync ingestion must be
        closed after run_until_complete (was never closed — leaked selector FD
        on every sync-context office edit)."""
        s, db, _ = svc
        p = tmp_path / "f.txt"
        p.write_bytes(b"content")

        class FakeLoop:
            closed = False

            def run_until_complete(self, coro):
                coro.close()
                return {"status": "ingested", "chars_ingested": 7}

            def close(self):
                self.closed = True

        with patch("core.auto_document_ingestion.AutoDocumentIngestionService") as cls, \
             patch("asyncio.new_event_loop", return_value=FakeLoop()) as nel:
            ing = cls.return_value
            ing.process_file_bytes = AsyncMock(return_value={"status": "ingested"})
            result = s._ingest_document_to_memory_sync(str(p), "u1")
        assert result is True
        assert nel.return_value.closed is True


class TestIngestAsync:
    def test_missing_file_returns_false(self, svc, tmp_path):
        s, db, _ = svc
        loop = asyncio.new_event_loop()
        try:
            result = loop.run_until_complete(
                s._ingest_document_to_memory(str(tmp_path / "nope"), "u1"))
        finally:
            loop.close()
        assert result is False

    def test_ingested_status_logs_and_returns_true(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "f.txt"
        p.write_bytes(b"content")
        with patch("core.auto_document_ingestion.AutoDocumentIngestionService") as cls:
            ing = cls.return_value
            ing.process_file_bytes = AsyncMock(
                return_value={"status": "ingested", "chars_ingested": 7})
            loop = asyncio.new_event_loop()
            try:
                result = loop.run_until_complete(
                    s._ingest_document_to_memory(str(p), "u1"))
            finally:
                loop.close()
        assert result is True

    def test_exception_returns_false(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "f.txt"
        p.write_bytes(b"x")
        with patch("core.auto_document_ingestion.AutoDocumentIngestionService",
                   side_effect=RuntimeError("boom")):
            loop = asyncio.new_event_loop()
            try:
                result = loop.run_until_complete(
                    s._ingest_document_to_memory(str(p), "u1"))
            finally:
                loop.close()
        assert result is False

    def test_skipped_status_returns_true(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "f.txt"
        p.write_bytes(b"content")
        with patch("core.auto_document_ingestion.AutoDocumentIngestionService") as cls:
            ing = cls.return_value
            ing.process_file_bytes = AsyncMock(return_value={"status": "skipped"})
            loop = asyncio.new_event_loop()
            try:
                result = loop.run_until_complete(
                    s._ingest_document_to_memory(str(p), "u1"))
            finally:
                loop.close()
        assert result is True

    def test_failed_status_returns_false(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "f.txt"
        p.write_bytes(b"content")
        with patch("core.auto_document_ingestion.AutoDocumentIngestionService") as cls:
            ing = cls.return_value
            ing.process_file_bytes = AsyncMock(return_value={"status": "failed"})
            loop = asyncio.new_event_loop()
            try:
                result = loop.run_until_complete(
                    s._ingest_document_to_memory(str(p), "u1"))
            finally:
                loop.close()
        assert result is False


class TestReadFileBytes:
    def test_empty_file_returns_none(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "empty.txt"
        p.write_bytes(b"")
        assert s._read_file_bytes(str(p)) is None

    def test_read_error_returns_none(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "locked.txt"
        p.write_bytes(b"x")
        with patch("builtins.open", side_effect=PermissionError("denied")):
            assert s._read_file_bytes(str(p)) is None
