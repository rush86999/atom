"""Coverage wave 61 — core/office_service.py (61% → 98%+).

Expands on wave 58 with the remaining branch/error paths: path-validation
edge cases (symlink escape, resolve OSError, base-itself), Excel read/write
(create-sheet, string cast variants, string-`=` cell type, no-coordinate,
save-failure, recalc success/failure/skip paths, all runtime delegations),
Word read/modify (blank paragraphs, tables, replace-without-target, unknown
action, corrupt file), PPTX read/modify (table shapes, layout clamp, unknown
action, library-unavailable), renderer (mammoth error, running-loop xlsx,
render exceptions, pptx unavailable) and manager dispatch case-insensitivity.
"""
import asyncio
import os
import tempfile
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock, Mock, patch

import openpyxl
import pytest

from core.office_service import (
    DocumentRenderer,
    ExcelManager,
    OfficeService,
    PowerPointManager,
    WordManager,
    _validate_office_path,
)


@pytest.fixture
def office_dir(tmp_path, monkeypatch):
    monkeypatch.setenv("ATOM_OFFICE_DIR", str(tmp_path))
    return tmp_path


def _make_xlsx(path, values=None):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws["A1"] = "hello"
    ws["A2"] = "world"
    ws["B1"] = "=1+1"
    wb.save(path)
    return path


def _make_docx(path, paragraphs=("original text",), table_cells=None):
    import docx
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_cells:
        t = doc.add_table(rows=1, cols=len(table_cells))
        for i, c in enumerate(table_cells):
            t.rows[0].cells[i].text = c
    doc.save(path)
    return path


class TestValidatePath:
    def test_empty_raises(self, office_dir):
        with pytest.raises(ValueError):
            _validate_office_path("")

    def test_valid_nested(self, office_dir):
        sub = office_dir / "nested"
        sub.mkdir()
        p = sub / "f.xlsx"
        p.write_bytes(b"")
        assert _validate_office_path(str(p)) == str(p.resolve())

    def test_base_itself_allowed(self, office_dir):
        assert _validate_office_path(str(office_dir)) == str(office_dir.resolve())

    def test_oserror_on_resolve_raises(self, office_dir):
        real_resolve = Path.resolve

        def fake_resolve(self, *a, **k):
            if str(self).startswith("/etc"):
                raise OSError("unresolvable")
            return real_resolve(self, *a, **k)

        with patch("core.office_service.Path.resolve", fake_resolve):
            with pytest.raises(ValueError):
                _validate_office_path("/etc/passwd")

    def test_symlink_escape_raises(self, office_dir):
        link = office_dir / "escape.xlsx"
        try:
            os.symlink("/etc/hosts", link)
        except OSError:
            pytest.skip("symlink not permitted on this host")
        with pytest.raises(ValueError):
            _validate_office_path(str(link))


class TestParsePath:
    def test_extra_variants(self):
        assert ExcelManager.parse_path("") == ("", "")
        assert ExcelManager.parse_path("  /Sheet1/A1  ") == ("Sheet1", "A1")
        assert ExcelManager.parse_path("//Sheet1//A1") == ("Sheet1", "A1")


class TestExcelRead:
    def test_invalid_path_error_dict(self, office_dir):
        result = ExcelManager().read_range("/etc/passwd", "A1")
        assert result["success"] is False

    def test_missing_file(self, office_dir):
        result = ExcelManager().read_range(str(office_dir / "nope.xlsx"), "A1")
        assert result["success"] is False
        assert "not found" in result["error"].lower()

    def test_overview_no_coordinate(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().read_range(str(p), "/Sheet1")
        assert result["success"] is True
        assert result["sheet_names"] == ["Sheet1"]
        assert result["dimensions"]

    def test_corrupt_file_generic_error(self, office_dir):
        p = office_dir / "bad.xlsx"
        p.write_bytes(b"not a real xlsx")
        result = ExcelManager().read_range(str(p), "A1")
        assert result["success"] is False
        assert result["error"] == "Failed to read Excel range"

    def test_single_cell_no_formula(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().read_range(str(p), "/Sheet1/A1")
        assert result["success"] is True
        assert result["value"] == "hello"
        assert result["formula"] is None
        assert result["cell_type"] == "text"

    def test_unknown_sheet_falls_back_to_active(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().read_range(str(p), "/NoSuchSheet/A1")
        assert result["success"] is True
        assert result["sheet_name"] == "Sheet1"

    def test_range_formula_cell_data_only_semantics(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().read_range(str(p), "/Sheet1/A1:B1")
        assert result["success"] is True
        # data_only=True workbook: unevaluated formula has no cached value
        assert result["cells"][0][1]["value"] is None
        assert result["cells"][0][1]["cell_type"] == "text"

    def test_range_without_sheet_prefix(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().read_range(str(p), "A1:A2")
        assert result["success"] is True
        assert len(result["cells"]) == 2


class TestExcelWrite:
    def test_invalid_path_error_dict(self, office_dir):
        result = ExcelManager().write_cell("/etc/passwd", "A1", 1)
        assert result["success"] is False

    def test_write_new_file(self, office_dir):
        p = office_dir / "new.xlsx"
        result = ExcelManager().write_cell(str(p), "/Sheet1/A1", 42)
        assert result["success"] is True
        assert openpyxl.load_workbook(p)["Sheet1"]["A1"].value == 42

    def test_write_default_sheet(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().write_cell(str(p), "A1", "no-sheet")
        assert result["success"] is True
        assert result["sheet_name"] == "Sheet1"

    def test_write_creates_missing_sheet(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().write_cell(str(p), "/NewSheet/B2", "x")
        assert result["success"] is True
        wb = openpyxl.load_workbook(p)
        assert "NewSheet" in wb.sheetnames
        assert wb["NewSheet"]["B2"].value == "x"

    def test_write_int_cast(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().write_cell(str(p), "/Sheet1/A1", "42")
        assert result["success"] is True
        assert openpyxl.load_workbook(p)["Sheet1"]["A1"].value == 42

    def test_write_float_cast(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().write_cell(str(p), "/Sheet1/A1", "4.5")
        assert result["success"] is True
        assert openpyxl.load_workbook(p)["Sheet1"]["A1"].value == 4.5

    def test_write_keeps_non_numeric_string(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().write_cell(str(p), "/Sheet1/A1", "abc123")
        assert result["success"] is True
        assert openpyxl.load_workbook(p)["Sheet1"]["A1"].value == "abc123"

    def test_formula_prefix_added(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().write_cell(str(p), "/Sheet1/C1", "SUM(A1:A2)",
                                           is_formula=True)
        assert result["success"] is True
        assert result["formula"] == "=SUM(A1:A2)"
        assert openpyxl.load_workbook(p)["Sheet1"]["C1"].value == "=SUM(A1:A2)"

    def test_string_equals_forced_text_type(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().write_cell(str(p), "/Sheet1/C1", "=not-a-formula")
        assert result["success"] is True
        assert openpyxl.load_workbook(p)["Sheet1"]["C1"].data_type == "s"

    def test_no_coordinate_error(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        result = ExcelManager().write_cell(str(p), "/Sheet1", 5)
        assert result["success"] is False
        assert "Cell coordinate" in result["error"]

    def test_save_failure_generic_error(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        with patch("core.office_service.openpyxl.Workbook.save",
                   side_effect=PermissionError("denied")):
            result = ExcelManager().write_cell(str(p), "/Sheet1/A1", 1)
        assert result["success"] is False
        assert result["error"] == "Failed to write Excel cell"

    def test_recalc_computed_value_returned(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        runtime = Mock()
        runtime.can_evaluate = True
        runtime.recalculate = AsyncMock(return_value=Path(p))

        class FakeCell:
            def __init__(self, value):
                self.value = value

        class FakeWS:
            def __getitem__(self, coord):
                return FakeCell(42)

        class FakeWB:
            sheetnames = ["Sheet1"]

            def __getitem__(self, name):
                return FakeWS()

        real_load = openpyxl.load_workbook

        def fake_load(path, data_only=False, **kw):
            if data_only:
                return FakeWB()
            return real_load(path, data_only=data_only, **kw)

        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime), \
             patch("core.office_service.openpyxl.load_workbook", fake_load):
            result = ExcelManager().write_cell(str(p), "/Sheet1/C1", "=1+1",
                                               is_formula=True)
        assert result["success"] is True
        assert result["value"] == 42
        runtime.recalculate.assert_awaited()

    def test_recalc_failure_keeps_value(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        runtime = Mock()
        runtime.can_evaluate = True
        runtime.recalculate = AsyncMock(side_effect=RuntimeError("soffice gone"))

        class FakeCell:
            def __init__(self, value):
                self.value = value

        class FakeWS:
            def __getitem__(self, coord):
                return FakeCell("=1+1")

        class FakeWB:
            sheetnames = ["Sheet1"]

            def __getitem__(self, name):
                return FakeWS()

        real_load = openpyxl.load_workbook

        def fake_load(path, data_only=False, **kw):
            if data_only:
                return FakeWB()
            return real_load(path, data_only=data_only, **kw)

        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime), \
             patch("core.office_service.openpyxl.load_workbook", fake_load):
            result = ExcelManager().write_cell(str(p), "/Sheet1/C1", "=1+1",
                                               is_formula=True)
        assert result["success"] is True
        assert result["value"] == "=1+1"

    def test_recalc_machinery_error_non_fatal(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        with patch("core.workbook_runtime.get_workbook_runtime",
                   side_effect=ImportError("no runtime")):
            result = ExcelManager().write_cell(str(p), "/Sheet1/A1", "plain")
        assert result["success"] is True
        assert result["value"] == "plain"

    def test_recalc_skipped_for_plain_value(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        runtime = Mock()
        runtime.can_evaluate = True
        runtime.recalculate = AsyncMock()
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime):
            result = ExcelManager().write_cell(str(p), "/Sheet1/A1", "plain")
        assert result["success"] is True
        runtime.recalculate.assert_not_awaited()


class TestExcelRuntimeDelegation:
    @pytest.fixture
    def runtime(self):
        rt = Mock()
        rt.insert_rows = AsyncMock(return_value={"success": True})
        rt.insert_cols = AsyncMock(return_value={"success": True})
        rt.get_evaluated_range = AsyncMock(return_value={"success": True, "cells": []})
        rt.recalculate = AsyncMock(return_value=Path("f.xlsx"))
        rt.add_pivot_table = AsyncMock(return_value={"success": True})
        rt.run_macro = AsyncMock(return_value={"success": True})
        rt.engine = "mock-engine"
        return rt

    @pytest.mark.asyncio
    async def test_insert_rows(self, office_dir, runtime):
        p = _make_xlsx(office_dir / "a.xlsx")
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime):
            res = await ExcelManager.insert_rows(str(p), "Sheet1", 2, 3)
        assert res == {"success": True}
        runtime.insert_rows.assert_awaited_once_with(str(p), "Sheet1", 2, 3)

    @pytest.mark.asyncio
    async def test_insert_columns(self, office_dir, runtime):
        p = _make_xlsx(office_dir / "a.xlsx")
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime):
            res = await ExcelManager.insert_columns(str(p), "Sheet1", 2, 3)
        assert res == {"success": True}
        runtime.insert_cols.assert_awaited_once_with(str(p), "Sheet1", 2, 3)

    @pytest.mark.asyncio
    async def test_get_evaluated_range(self, office_dir, runtime):
        p = _make_xlsx(office_dir / "a.xlsx")
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime):
            res = await ExcelManager.get_evaluated_range(str(p), "/Sheet1/A1:B2")
        assert res["success"] is True
        runtime.get_evaluated_range.assert_awaited_once_with(
            str(p), "Sheet1", "A1", "B2")

    @pytest.mark.asyncio
    async def test_get_evaluated_range_single_cell(self, office_dir, runtime):
        p = _make_xlsx(office_dir / "a.xlsx")
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime):
            res = await ExcelManager.get_evaluated_range(str(p), "/Sheet1/A1")
        assert res["success"] is True
        runtime.get_evaluated_range.assert_awaited_once_with(
            str(p), "Sheet1", "A1", None)

    @pytest.mark.asyncio
    async def test_recalculate(self, office_dir, runtime):
        p = _make_xlsx(office_dir / "a.xlsx")
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime):
            res = await ExcelManager.recalculate(str(p))
        assert res == {"success": True, "engine": "mock-engine"}
        runtime.recalculate.assert_awaited_once()

    @pytest.mark.asyncio
    async def test_add_pivot_table(self, office_dir, runtime):
        p = _make_xlsx(office_dir / "a.xlsx")
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime):
            res = await ExcelManager.add_pivot_table(
                str(p), "Sheet1", "Pivot", "A1:B2", ["A"], ["B"], [{"col": "A"}])
        assert res == {"success": True}
        runtime.add_pivot_table.assert_awaited_once()

    @pytest.mark.asyncio
    async def test_run_excel_macro(self, office_dir, runtime):
        p = _make_xlsx(office_dir / "a.xlsx")
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime):
            res = await ExcelManager.run_excel_macro(str(p), "MyMacro")
        assert res == {"success": True}
        runtime.run_macro.assert_awaited_once_with(str(p), "MyMacro")

    @pytest.mark.asyncio
    async def test_all_missing_file(self, office_dir, runtime):
        p = str(office_dir / "nope.xlsx")
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime):
            assert (await ExcelManager.insert_rows(p, "S", 1))["success"] is False
            assert (await ExcelManager.insert_columns(p, "S", 1))["success"] is False
            assert (await ExcelManager.get_evaluated_range(p, "A1"))["success"] is False
            assert (await ExcelManager.recalculate(p))["success"] is False
            assert (await ExcelManager.add_pivot_table(
                p, "S", "P", "A1", [], [], []))["success"] is False
            assert (await ExcelManager.run_excel_macro(p, "M"))["success"] is False

    @pytest.mark.asyncio
    async def test_all_invalid_path(self, office_dir, runtime):
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime):
            assert (await ExcelManager.insert_rows("/etc/passwd", "S", 1))["success"] is False
            assert (await ExcelManager.insert_columns("/etc/passwd", "S", 1))["success"] is False
            assert (await ExcelManager.get_evaluated_range("/etc/passwd", "A1"))["success"] is False
            assert (await ExcelManager.recalculate("/etc/passwd"))["success"] is False
            assert (await ExcelManager.add_pivot_table(
                "/etc/passwd", "S", "P", "A1", [], [], []))["success"] is False
            assert (await ExcelManager.run_excel_macro("/etc/passwd", "M"))["success"] is False


class TestWord:
    def test_invalid_path(self, office_dir):
        result = WordManager().read_document("/etc/passwd")
        assert result["success"] is False

    def test_missing_file(self, office_dir):
        result = WordManager().read_document(str(office_dir / "nope.docx"))
        assert result["success"] is False
        assert "not found" in result["error"].lower()

    def test_modify_replace_paragraph(self, office_dir):
        p = _make_docx(office_dir / "d.docx", paragraphs=("OLD placeholder",))
        result = WordManager().modify_document(
            str(p), "replace", "NEW", options={"target": "OLD"})
        assert result["success"] is True
        import docx
        assert docx.Document(p).paragraphs[0].text == "NEW placeholder"

    def test_read_skips_blank_paragraphs_and_tables(self, office_dir):
        p = _make_docx(office_dir / "d.docx",
                       paragraphs=("  ", "real text"),
                       table_cells=("cell-a", ""))
        result = WordManager().read_document(str(p))
        assert result["success"] is True
        assert [x["text"] for x in result["paragraphs"]] == ["real text"]
        assert result["metadata"]["tables_count"] == 1
        assert result["tables"][0]["rows"][0] == ["cell-a", ""]

    def test_modify_replace_in_table(self, office_dir):
        p = _make_docx(office_dir / "d.docx",
                       paragraphs=("keep this",),
                       table_cells=("OLD target", "x"))
        result = WordManager().modify_document(
            str(p), "replace", "NEW", options={"target": "OLD"})
        assert result["success"] is True
        import docx
        read = docx.Document(p)
        assert read.tables[0].rows[0].cells[0].text == "NEW target"

    def test_replace_without_target_error(self, office_dir):
        p = _make_docx(office_dir / "d.docx")
        result = WordManager().modify_document(str(p), "replace", "x", options={})
        assert result["success"] is False
        assert "requires a target" in result["error"]

    def test_unknown_action_error(self, office_dir):
        p = _make_docx(office_dir / "d.docx")
        result = WordManager().modify_document(str(p), "delete", "x")
        assert result["success"] is False
        assert "Unknown modification action" in result["error"]

    def test_modify_creates_new_file(self, office_dir):
        p = office_dir / "new.docx"
        result = WordManager().modify_document(str(p), "append", "first")
        assert result["success"] is True
        import docx
        assert docx.Document(p).paragraphs[0].text == "first"

    def test_modify_corrupt_error(self, office_dir):
        p = office_dir / "bad.docx"
        p.write_bytes(b"not a docx")
        result = WordManager().modify_document(str(p), "append", "x")
        assert result["success"] is False

    def test_read_corrupt_error(self, office_dir):
        p = office_dir / "bad.docx"
        p.write_bytes(b"not a docx")
        result = WordManager().read_document(str(p))
        assert result["success"] is False

    def test_modify_invalid_path(self, office_dir):
        result = WordManager().modify_document("/etc/passwd", "append", "x")
        assert result["success"] is False


class FakePptx:
    """Deterministic python-pptx stand-in: text + table shapes, layouts, slides."""

    def __init__(self, n_layouts=3, slides=()):
        self.slide_layouts = [SimpleNamespace() for _ in range(n_layouts)]
        self.slides = FakeSlides(slides)
        self.save = Mock()

    @staticmethod
    def _new_prs():
        return FakePptx()

    @staticmethod
    def _mock_module():
        fake_mod = Mock()
        fake_mod.Presentation.side_effect = FakePptx._new_prs
        return fake_mod


class FakeCell:
    def __init__(self, text):
        self.text = text


class FakeRow:
    def __init__(self, texts):
        self.cells = [FakeCell(t) for t in texts]


class FakeTable:
    def __init__(self, rows):
        self.rows = [FakeRow(r) for r in rows]


class FakeShape:
    def __init__(self, name, kind, text="", table_rows=None):
        self.name = name
        self.has_text_frame = kind == "text"
        self.has_table = kind == "table"
        self.text_frame = SimpleNamespace(text=text) if self.has_text_frame else None
        self.table = FakeTable(table_rows or []) if self.has_table else None


class FakeShapes:
    def __init__(self, shapes=()):
        self._items = list(shapes)
        self.title = SimpleNamespace(text="")

    def __iter__(self):
        return iter(self._items)


class FakeSlide:
    def __init__(self, shapes=()):
        self.shapes = FakeShapes(shapes)
        self.placeholders = [SimpleNamespace(text=""), SimpleNamespace(text="")]


class FakeSlides:
    def __init__(self, slides=()):
        self._items = list(slides)

    def __iter__(self):
        return iter(self._items)

    def __len__(self):
        return len(self._items)

    def add_slide(self, layout):
        new = FakeSlide()
        self._items.append(new)
        return new


@pytest.fixture
def fake_pptx():
    import core.office_service as osvc
    deck = FakePptx(slides=[
        FakeSlide([
            FakeShape("t1", "text", text="Slide One Text"),
            FakeShape("tbl", "table", table_rows=[["a1", "b1"], ["a2", "b2"]]),
        ])
    ])
    fake_mod = Mock()
    fake_mod.Presentation.return_value = deck
    with patch.object(osvc, "pptx", fake_mod, create=True), \
         patch.object(osvc, "PPTX_AVAILABLE", True):
        yield deck


class TestPptx:
    def test_invalid_path_read(self, office_dir, fake_pptx):
        result = PowerPointManager().read_slides("/etc/passwd")
        assert result["success"] is False

    def test_missing_file_read(self, office_dir, fake_pptx):
        result = PowerPointManager().read_slides(str(office_dir / "nope.pptx"))
        assert result["success"] is False
        assert "not found" in result["error"].lower()

    def test_read_text_and_table_shapes(self, office_dir, fake_pptx):
        p = office_dir / "d.pptx"
        p.write_bytes(b"x")
        result = PowerPointManager().read_slides(str(p))
        assert result["success"] is True
        assert result["slide_count"] == 1
        shapes = result["slides"][0]["shapes"]
        assert shapes[0] == {"type": "text", "name": "t1", "text": "Slide One Text"}
        assert shapes[1]["type"] == "table"
        assert shapes[1]["table"] == [["a1", "b1"], ["a2", "b2"]]

    def test_not_available_error(self, office_dir):
        import core.office_service as osvc
        p = office_dir / "d.pptx"
        p.write_bytes(b"x")
        with patch.object(osvc, "PPTX_AVAILABLE", False):
            result = PowerPointManager().read_slides(str(p))
        assert result["success"] is False
        assert "not installed" in result["error"]

    def test_read_corrupt_error(self, office_dir, fake_pptx):
        import core.office_service as osvc
        p = office_dir / "d.pptx"
        p.write_bytes(b"not a pptx")
        osvc.pptx.Presentation.side_effect = RuntimeError("bad zip")
        result = PowerPointManager().read_slides(str(p))
        assert result["success"] is False
        assert "Failed to read PowerPoint" in result["error"]

    def test_modify_add_slide_full(self, office_dir, fake_pptx):
        p = office_dir / "d.pptx"
        p.write_bytes(b"x")
        result = PowerPointManager().modify_slides(
            str(p), "add_slide",
            {"title": "New Title", "content": "Body", "layout_idx": 2})
        assert result["success"] is True
        new_slide = fake_pptx.slides._items[-1]
        assert new_slide.shapes.title.text == "New Title"
        assert new_slide.placeholders[1].text == "Body"

    def test_modify_layout_clamped(self, office_dir, fake_pptx):
        p = office_dir / "d.pptx"
        p.write_bytes(b"x")
        result = PowerPointManager().modify_slides(
            str(p), "add_slide", {"title": None, "content": None, "layout_idx": 99})
        assert result["success"] is True
        # clamped to index 1 — deck has 3 layouts so 1 stays in range
        assert len(fake_pptx.slides._items) == 2

    def test_modify_unknown_action(self, office_dir, fake_pptx):
        p = office_dir / "d.pptx"
        p.write_bytes(b"x")
        result = PowerPointManager().modify_slides(str(p), "delete", {})
        assert result["success"] is False
        assert "Unknown PowerPoint action" in result["error"]

    def test_modify_not_available_error(self, office_dir):
        import core.office_service as osvc
        p = office_dir / "d.pptx"
        p.write_bytes(b"x")
        with patch.object(osvc, "PPTX_AVAILABLE", False):
            result = PowerPointManager().modify_slides(str(p), "add_slide", {})
        assert result["success"] is False

    def test_modify_corrupt_error(self, office_dir, fake_pptx):
        import core.office_service as osvc
        p = office_dir / "d.pptx"
        p.write_bytes(b"not a pptx")
        osvc.pptx.Presentation.side_effect = RuntimeError("bad zip")
        result = PowerPointManager().modify_slides(str(p), "add_slide", {})
        assert result["success"] is False

    def test_modify_invalid_path(self, office_dir, fake_pptx):
        result = PowerPointManager().modify_slides("/etc/passwd", "add_slide", {})
        assert result["success"] is False


class TestRenderer:
    def test_invalid_path(self, office_dir):
        result = DocumentRenderer.render_to_html("/etc/passwd")
        assert result["success"] is False

    def test_docx_render_mammoth_missing(self, office_dir):
        import docx
        p = office_dir / "d.docx"
        doc = docx.Document()
        doc.add_paragraph("hello")
        doc.save(p)
        import core.office_service as osvc
        with patch.object(osvc, "MAMMOTH_AVAILABLE", False):
            result = DocumentRenderer.render_to_html(str(p))
        assert result["success"] is False
        assert "mammoth library not installed" in result["error"]

    def test_docx_render_mammoth_error(self, office_dir):
        import docx
        p = office_dir / "d.docx"
        doc = docx.Document()
        doc.add_paragraph("hello")
        doc.save(p)
        mammoth = Mock()
        mammoth.convert_to_html.side_effect = RuntimeError("convert failed")
        with patch("core.office_service.mammoth", mammoth, create=True), \
             patch("core.office_service.MAMMOTH_AVAILABLE", True):
            result = DocumentRenderer.render_to_html(str(p))
        assert result["success"] is False
        assert "Failed rendering Word to HTML" in result["error"]

    def test_unsupported_format(self, office_dir):
        p = office_dir / "f.txt"
        p.write_bytes(b"x")
        result = DocumentRenderer.render_to_html(str(p))
        assert result["success"] is False
        assert "Unsupported" in result["error"]

    def test_docx_render_with_warnings(self, office_dir):
        import docx
        p = office_dir / "d.docx"
        doc = docx.Document()
        doc.add_paragraph("hello")
        doc.save(p)
        mammoth = Mock()
        mammoth.convert_to_html.return_value = SimpleNamespace(
            value="<p>hello</p>", messages=[SimpleNamespace(message="warn1")])
        with patch("core.office_service.mammoth", mammoth, create=True), \
             patch("core.office_service.MAMMOTH_AVAILABLE", True):
            result = DocumentRenderer.render_to_html(str(p))
        assert result["success"] is True
        assert result["warnings"] == ["warn1"]

    def test_xlsx_render_running_loop_basic(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        runtime = Mock()
        runtime.engine = "basic"
        runtime._render_html_basic.return_value = "<table>basic</table>"
        loop = Mock()
        loop.is_running.return_value = True
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime), \
             patch("asyncio.get_event_loop", return_value=loop):
            result = DocumentRenderer.render_to_html(str(p))
        assert result["success"] is True
        runtime._render_html_basic.assert_called_once()
        loop.run_until_complete.assert_not_called()

    def test_xlsx_render_idle_loop_await(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        runtime = Mock()
        runtime.engine = "soffice"
        runtime.render_to_html = AsyncMock(return_value="<table>soffice</table>")
        loop = Mock()
        loop.is_running.return_value = False
        loop.run_until_complete.return_value = "<table>soffice</table>"
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime), \
             patch("asyncio.get_event_loop", return_value=loop):
            result = DocumentRenderer.render_to_html(str(p))
        assert result["success"] is True
        assert result["engine"] == "soffice"
        loop.run_until_complete.assert_called_once()

    def test_xlsx_render_no_loop_basic(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        runtime = Mock()
        runtime.engine = "basic"
        runtime._render_html_basic.return_value = "<table>basic</table>"
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime), \
             patch("asyncio.get_event_loop",
                   side_effect=RuntimeError("no current loop")):
            result = DocumentRenderer.render_to_html(str(p))
        assert result["success"] is True
        runtime._render_html_basic.assert_called_once()

    def test_xlsx_render_exception_generic(self, office_dir):
        p = _make_xlsx(office_dir / "a.xlsx")
        runtime = Mock()
        runtime.engine = "basic"
        runtime._render_html_basic.side_effect = RuntimeError("boom")
        with patch("core.workbook_runtime.get_workbook_runtime",
                   return_value=runtime), \
             patch("asyncio.get_event_loop",
                   side_effect=RuntimeError("no current loop")):
            result = DocumentRenderer.render_to_html(str(p))
        assert result["success"] is False
        assert "Failed rendering Excel to HTML" in result["error"]

    def test_pptx_render_text_only(self, office_dir, fake_pptx):
        import core.office_service as osvc
        p = office_dir / "d.pptx"
        p.write_bytes(b"x")
        deck = FakePptx(slides=[FakeSlide([FakeShape("t1", "text", text="Hi")])])
        osvc.pptx.Presentation.return_value = deck
        result = DocumentRenderer.render_to_html(str(p))
        assert result["success"] is True
        assert "<p>Hi</p>" in result["html"]
        assert "Slide 1" in result["html"]

    def test_pptx_render_not_available(self, office_dir):
        import core.office_service as osvc
        p = office_dir / "d.pptx"
        p.write_bytes(b"x")
        with patch.object(osvc, "PPTX_AVAILABLE", False):
            result = DocumentRenderer.render_to_html(str(p))
        assert result["success"] is False
        assert "not installed" in result["error"]

    def test_pptx_render_error(self, office_dir, fake_pptx):
        import core.office_service as osvc
        p = office_dir / "d.pptx"
        p.write_bytes(b"x")
        osvc.pptx.Presentation.side_effect = RuntimeError("bad zip")
        result = DocumentRenderer.render_to_html(str(p))
        assert result["success"] is False
        assert "Failed rendering PPTX to HTML" in result["error"]


class TestService:
    def test_manager_dispatch_case_insensitive(self):
        svc = OfficeService()
        assert isinstance(svc.get_manager_for_file("A.XLSX"), ExcelManager)
        assert isinstance(svc.get_manager_for_file("B.Docx"), WordManager)
        assert isinstance(svc.get_manager_for_file("C.PPTX"), PowerPointManager)
        with pytest.raises(ValueError):
            svc.get_manager_for_file("a.txt")

    def test_components_initialized(self):
        svc = OfficeService()
        assert isinstance(svc.excel, ExcelManager)
        assert isinstance(svc.word, WordManager)
        assert isinstance(svc.pptx, PowerPointManager)
        assert isinstance(svc.renderer, DocumentRenderer)
