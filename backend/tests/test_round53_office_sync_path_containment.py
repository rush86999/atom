"""
Round 53 — Office sync: missing path containment (arbitrary file read/overwrite)
(Red-Green-Refactor).

office_routes /sync-update and /present pass user-supplied file_path into
core/office_sync_service, which NEVER validates the path against the office
directory (unlike every OfficeService entry point):

  A. sync_canvas_to_file() — docx branch does `doc.save(file_path)` on an
     unvalidated path: any AUTHENTICATED user can overwrite any existing
     .docx file the process can write (arbitrary file modification).

  B. broadcast_file_update() — renders and reads (and ingests to memory) any
     existing office file the process can read; the rendered HTML is pushed
     to the caller's own canvas_id WebSocket + CanvasAudit rows (arbitrary
     file-read exfiltration).

  C. sync_canvas_to_file() also leaks `str(e)` in its failure dict (R52 class).

Fix: validate file_path with office_service._validate_office_path() at both
sync entry points; generic error string on the exception path.
"""

from datetime import datetime
from unittest.mock import MagicMock, patch

import pytest

from fastapi import FastAPI
from fastapi.testclient import TestClient

from core.auth import get_current_user as auth_get_current_user
from core.database import get_db_session

SECRET = "secret-sync-internal-xyz"


def _make_client(monkeypatch):
    from api.office_routes import router

    app = FastAPI()
    app.include_router(router)
    app.dependency_overrides[auth_get_current_user] = lambda: MagicMock(
        id="u-53", email="u@example.com"
    )
    app.dependency_overrides[get_db_session] = lambda: MagicMock()
    return TestClient(app, raise_server_exceptions=False)


class TestSyncPathContainment:
    def test_docx_sync_outside_dir_rejected_without_write(self, monkeypatch, tmp_path):
        """Existing .docx OUTSIDE the office dir must not be overwritten."""
        monkeypatch.setenv("ATOM_OFFICE_DIR", str(tmp_path / "office"))
        outside = tmp_path / "outside" / "victim.docx"
        outside.parent.mkdir()
        import docx

        doc = docx.Document()
        doc.add_paragraph("ORIGINAL CONTENT")
        doc.save(outside)
        original_bytes = outside.read_bytes()

        from core.office_sync_service import OfficeSyncService

        res = OfficeSyncService(MagicMock()).sync_canvas_to_file(
            canvas_id="c1",
            file_path=str(outside),
            user_id="u1",
            edit_type="document",
            data={"content": "ATTACKER OVERWRITE"},
        )

        assert res.get("success") is False, (
            "sync_canvas_to_file overwrote a file outside the office dir — "
            "arbitrary file modification"
        )
        assert outside.read_bytes() == original_bytes, (
            "victim .docx was modified by an out-of-scope sync"
        )

    def test_excel_sync_outside_dir_rejected(self, monkeypatch, tmp_path):
        monkeypatch.setenv("ATOM_OFFICE_DIR", str(tmp_path / "office"))
        outside = tmp_path / "outside" / "book.xlsx"
        outside.parent.mkdir()
        outside.write_bytes(b"not really xlsx but existing")

        from core.office_sync_service import OfficeSyncService

        res = OfficeSyncService(MagicMock()).sync_canvas_to_file(
            canvas_id="c1",
            file_path=str(outside),
            user_id="u1",
            edit_type="cell",
            data={"cell_path": "/Sheet1/A1", "value": "EVIL"},
        )

        assert res.get("success") is False

    def test_broadcast_outside_dir_does_not_read_or_audit(self, monkeypatch, tmp_path):
        """broadcast_file_update must not render/read/ingest out-of-scope files."""
        monkeypatch.setenv("ATOM_OFFICE_DIR", str(tmp_path / "office"))
        outside = tmp_path / "outside" / "doc.docx"
        outside.parent.mkdir()
        import docx

        doc = docx.Document()
        doc.add_paragraph("SENSITIVE")
        doc.save(outside)

        db = MagicMock()
        from core.office_sync_service import OfficeSyncService

        # Patch the ingestion + renderer so any attempt to touch the file is visible.
        with patch.object(
            OfficeSyncService, "_ingest_document_to_memory_sync"
        ) as ingest, patch(
            "core.office_service.DocumentRenderer.render_to_html"
        ) as render:
            OfficeSyncService(db).broadcast_file_update("c1", str(outside), "u1")

        render.assert_not_called()
        ingest.assert_not_called()
        db.add.assert_not_called()

    def test_sync_error_does_not_leak(self, monkeypatch, tmp_path):
        """In-scope sync failure must not leak the exception string (R52 class)."""
        monkeypatch.setenv("ATOM_OFFICE_DIR", str(tmp_path / "office"))
        in_scope = tmp_path / "office" / "book.xlsx"
        in_scope.parent.mkdir()
        in_scope.write_bytes(b"dummy")

        from core.office_sync_service import OfficeSyncService

        svc = OfficeSyncService(MagicMock())
        with patch.object(
            svc.office.excel, "write_cell", side_effect=RuntimeError(SECRET)
        ):
            res = svc.sync_canvas_to_file(
                canvas_id="c1",
                file_path=str(in_scope),
                user_id="u1",
                edit_type="cell",
                data={"cell_path": "/Sheet1/A1", "value": 1},
            )

        assert res.get("success") is False
        assert SECRET not in res.get("error", ""), (
            f"sync_canvas_to_file leaks internal exception detail: {res.get('error')!r}"
        )

    def test_http_sync_update_outside_dir_rejected(self, monkeypatch, tmp_path):
        """POST /sync-update with an out-of-scope path must fail closed."""
        monkeypatch.setenv("ATOM_OFFICE_DIR", str(tmp_path / "office"))
        outside = tmp_path / "outside" / "doc.docx"
        outside.parent.mkdir()
        import docx

        doc = docx.Document()
        doc.add_paragraph("ORIGINAL")
        doc.save(outside)
        original_bytes = outside.read_bytes()

        client = _make_client(monkeypatch)
        resp = client.post(
            "/sync-update",
            json={
                "canvas_id": "c1",
                "file_path": str(outside),
                "user_id": "u1",
                "edit_type": "document",
                "data": {"content": "ATTACKER"},
            },
        )

        assert resp.status_code == 400, (
            f"sync-update accepted an out-of-scope file_path ({resp.status_code})"
        )
        assert outside.read_bytes() == original_bytes, (
            "victim .docx was modified through the HTTP surface"
        )
