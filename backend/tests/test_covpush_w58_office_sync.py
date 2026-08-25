"""Coverage wave 58 — core/office_sync_service.py (37% → 90%+).

sync_canvas_to_file (path containment, missing file, xlsx cell write,
docx content rewrite, unsupported type, exception), broadcast_file_update
(render fail, audit row, async ingest + no-loop fallbacks, ws broadcast),
ingest async/sync paths, read-file bytes, validation rejects.
"""
import asyncio
import os
import tempfile
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock, Mock, patch

import openpyxl
import pytest

from core.office_sync_service import OfficeSyncService


@pytest.fixture
def svc(tmp_path, monkeypatch):
    monkeypatch.setenv("ATOM_OFFICE_DIR", str(tmp_path))
    db = Mock()
    return OfficeSyncService(db), db, tmp_path


class TestSyncCanvasToFile:
    def test_path_containment(self, svc):
        s, db, _ = svc
        result = s.sync_canvas_to_file("c1", "/etc/passwd", "u1", "cell", {})
        assert result["success"] is False

    def test_missing_file(self, svc):
        s, db, tmp = svc
        result = s.sync_canvas_to_file("c1", str(tmp / "missing.xlsx"), "u1", "cell", {})
        assert result["success"] is False
        assert "not found" in result["error"]

    def test_xlsx_cell_success(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "a.xlsx"
        wb = openpyxl.Workbook()
        wb.save(p)
        with patch.object(s.office.excel, "write_cell",
                          return_value={"success": True}), \
             patch.object(s, "broadcast_file_update") as bcast:
            result = s.sync_canvas_to_file(
                "c1", str(p), "u1", "cell", {"cell_path": "/Sheet1/A1", "value": 5})
        assert result["success"] is True
        bcast.assert_called_once()

    def test_xlsx_missing_cell_path(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "a.xlsx"
        wb = openpyxl.Workbook()
        wb.save(p)
        result = s.sync_canvas_to_file("c1", str(p), "u1", "cell", {})
        assert result["success"] is False
        assert "cell_path required" in result["error"]

    def test_xlsx_write_failure_propagates(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "a.xlsx"
        wb = openpyxl.Workbook()
        wb.save(p)
        with patch.object(s.office.excel, "write_cell",
                          return_value={"success": False, "error": "bad"}):
            result = s.sync_canvas_to_file(
                "c1", str(p), "u1", "cell", {"cell_path": "/S/A1", "value": 1})
        assert result["success"] is False

    def test_docx_content_rewrite(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "d.docx"
        import docx
        doc = docx.Document()
        doc.add_paragraph("old")
        doc.save(p)
        with patch.object(s, "broadcast_file_update") as bcast:
            result = s.sync_canvas_to_file(
                "c1", str(p), "u1", "document", {"content": "line1\nline2"})
        assert result["success"] is True
        bcast.assert_called_once()
        read = docx.Document(p)
        assert [para.text for para in read.paragraphs] == ["line1", "line2"]

    def test_unsupported_type(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "a.txt"
        p.write_bytes(b"x")
        result = s.sync_canvas_to_file("c1", str(p), "u1", "cell", {})
        assert result["success"] is False

    def test_exception_tolerated(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "a.xlsx"
        p.write_bytes(b"not xlsx")
        with patch.object(s.office.excel, "write_cell",
                          side_effect=RuntimeError("boom")):
            result = s.sync_canvas_to_file(
                "c1", str(p), "u1", "cell", {"cell_path": "/S/A1", "value": 1})
        assert result["success"] is False


class TestBroadcast:
    def test_invalid_path_returns(self, svc):
        s, db, _ = svc
        s.broadcast_file_update("c1", "/etc/passwd", "u1")  # must not raise

    def test_render_failure_returns(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "a.xlsx"
        wb = openpyxl.Workbook()
        wb.save(p)
        with patch.object(s.office.renderer, "render_to_html",
                          return_value={"success": False}):
            s.broadcast_file_update("c1", str(p), "u1")
        # Render failure degrades to html=None — the audit row is still
        # written (structured snapshot is independent of the render).
        audit = db.add.call_args[0][0]
        assert audit.details_json["html"] is None

    def test_success_with_audit_and_ws(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "a.xlsx"
        wb = openpyxl.Workbook()
        wb.save(p)
        with patch.object(s.office.renderer, "render_to_html",
                          return_value={"success": True, "html": "<b>x</b>"}), \
             patch.object(s, "_ingest_document_to_memory",
                          new=AsyncMock(return_value=True)), \
             patch("core.office_sync_service.ws_manager.broadcast",
                   new=AsyncMock()) as wsb:
            s.broadcast_file_update("c1", str(p), "u1")
        db.add.assert_called_once()
        db.commit.assert_called_once()
        wsb.assert_called()

    def test_no_loop_falls_back_to_sync(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "a.xlsx"
        wb = openpyxl.Workbook()
        wb.save(p)
        with patch.object(s.office.renderer, "render_to_html",
                          return_value={"success": True, "html": "<b>x</b>"}), \
             patch.object(s, "_ingest_document_to_memory",
                          new=AsyncMock(return_value=True)), \
             patch.object(s, "_ingest_document_to_memory_sync",
                          return_value=True) as sync_ingest, \
             patch("core.office_sync_service.ws_manager.broadcast",
                   new=AsyncMock()), \
             patch("asyncio.create_task", side_effect=RuntimeError("no loop")):
            s.broadcast_file_update("c1", str(p), "u1")
        sync_ingest.assert_called_once()


class TestIngest:
    def test_read_file_bytes(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "f.txt"
        p.write_bytes(b"content")
        assert s._read_file_bytes(str(p)) == b"content"
        assert s._read_file_bytes(str(tmp_path / "missing")) is None

    def test_ingest_async_success(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "f.txt"
        p.write_bytes(b"content")
        with patch("core.auto_document_ingestion.AutoDocumentIngestionService") as cls:
            ing = cls.return_value
            ing.process_file_bytes = AsyncMock(
                return_value={"status": "ingested", "chars_ingested": 7})
            loop = asyncio.new_event_loop()
            try:
                result = loop.run_until_complete(
                    s._ingest_document_to_memory(str(p), "u1"))
            finally:
                loop.close()
        assert result is True

    def test_ingest_async_missing_file(self, svc, tmp_path):
        s, db, _ = svc
        loop = asyncio.new_event_loop()
        try:
            result = loop.run_until_complete(
                s._ingest_document_to_memory(str(tmp_path / "nope"), "u1"))
        finally:
            loop.close()
        assert result is False

    def test_ingest_async_exception(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "f.txt"
        p.write_bytes(b"x")
        with patch("core.auto_document_ingestion.AutoDocumentIngestionService",
                   side_effect=RuntimeError("boom")):
            loop = asyncio.new_event_loop()
            try:
                result = loop.run_until_complete(
                    s._ingest_document_to_memory(str(p), "u1"))
            finally:
                loop.close()
        assert result is False

    def test_ingest_sync_success(self, svc, tmp_path):
        s, db, _ = svc
        p = tmp_path / "f.txt"
        p.write_bytes(b"content")
        with patch("core.auto_document_ingestion.AutoDocumentIngestionService") as cls:
            ing = cls.return_value
            ing.process_file_bytes = AsyncMock(
                return_value={"status": "skipped"})
            result = s._ingest_document_to_memory_sync(str(p), "u1")
        assert result is True
