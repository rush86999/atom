"""
Automatic Document Ingestion Service for Atom Memory
Auto-ingests documents from connected file storage integrations (Google Drive, Dropbox, etc.)
Supports: Excel, PDF, DOC/DOCX, TXT, CSV, Markdown files
"""

import asyncio
from dataclasses import dataclass, field
from datetime import datetime, timedelta, timezone
from enum import Enum
import io
import json
import logging
import os
from typing import Any, Dict, List, Optional, Set

# Import for lazy loading to avoid circular imports
# from core.atom_meta_agent import handle_data_event_trigger

logger = logging.getLogger(__name__)

# Strong refs for fire-and-forget post-ingestion agent-trigger tasks (a bare
# create_task result can be garbage-collected mid-flight).
_pending_agent_trigger_tasks: set = set()


class FileType(str, Enum):
    """Supported file types for ingestion"""
    PDF = "pdf"
    DOC = "doc"
    DOCX = "docx"
    TXT = "txt"
    CSV = "csv"
    EXCEL = "xlsx"
    XLS = "xls"
    MARKDOWN = "md"
    JSON = "json"


class IntegrationSource(str, Enum):
    """Supported file storage integrations"""
    GOOGLE_DRIVE = "google_drive"
    DROPBOX = "dropbox"
    ONEDRIVE = "onedrive"
    BOX = "box"
    SHAREPOINT = "sharepoint"
    NOTION = "notion"
    LOCAL = "local"


@dataclass
class IngestionSettings:
    """Settings for document ingestion per integration"""
    integration_id: str
    workspace_id: str
    enabled: bool = False
    auto_sync_new_files: bool = True
    file_types: List[str] = field(default_factory=lambda: ["pdf", "docx", "txt", "md"])
    sync_folders: List[str] = field(default_factory=list)  # Empty = all folders
    exclude_folders: List[str] = field(default_factory=list)
    max_file_size_mb: int = 50
    sync_frequency_minutes: int = 60
    last_sync: Optional[datetime] = None


@dataclass
class IngestedDocument:
    """Record of an ingested document"""
    id: str
    file_name: str
    file_path: str
    file_type: str
    integration_id: str
    workspace_id: str
    file_size_bytes: int
    content_preview: str  # First 500 chars
    ingested_at: datetime
    external_id: str  # ID in the source system
    external_modified_at: Optional[datetime] = None
    # Freshness tracking (mirrors the ORM columns added in core/models.py).
    # See core/doc_freshness_service.py for semantics.
    source_url: Optional[str] = None
    source_content_hash: Optional[str] = None
    last_verified_at: Optional[datetime] = None
    source_modified_at: Optional[datetime] = None
    freshness_status: str = "fresh"  # fresh|stale|outdated|removed|superseded
    superseded_by: Optional[str] = None  # id of a newer same-topic doc


def _parse_source_modified(raw: Any) -> Optional[datetime]:
    """Parse a connector-reported modified_at into an aware UTC datetime.

    Sources use wildly different shapes: WorkDrive listings render
    "May 1, 2025, 12:16 PM", APIs return ISO strings, some callers pass a
    datetime already. None → None (unknown, comparisons must not guess).
    """
    if isinstance(raw, datetime):
        return raw if raw.tzinfo else raw.replace(tzinfo=timezone.utc)
    if not isinstance(raw, str) or not raw.strip():
        return None
    raw = raw.strip()
    for fmt in (
        "%b %d, %Y, %I:%M %p",  # May 1, 2025, 12:16 PM (WorkDrive listings)
        "%b %d, %Y",  # May 1, 2025
        "%Y-%m-%dT%H:%M:%S.%f%z",  # ISO with micros + offset
        "%Y-%m-%dT%H:%M:%S%z",  # ISO + offset
        "%Y-%m-%dT%H:%M:%S.%f",  # ISO naive
        "%Y-%m-%dT%H:%M:%S",  # ISO naive
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d",
    ):
        try:
            dt = datetime.strptime(raw, fmt)
            return dt if dt.tzinfo else dt.replace(tzinfo=timezone.utc)
        except ValueError:
            continue
    return None


def _stored_copy_older_than(stored: Optional[Dict[str, Any]], source_modified_dt: Optional[datetime]) -> bool:
    """True when a stored document copy is KNOWN to predate the source's
    current modified_at. Baseline precedence: the modified time captured at
    last ingest, else the ingest/creation time itself (a source modified
    after we ingested is newer by definition). Conservative: when NO
    baseline and no source time can be established → False, because 'can't
    tell' must not trigger a re-ingest."""
    if not stored or source_modified_dt is None:
        return False
    meta = stored.get("metadata") or {}
    for raw in (
        meta.get("source_modified_at"),
        stored.get("source_modified_at"),  # top-level freshness column
        meta.get("external_modified_at"),
        stored.get("external_modified_at"),
        meta.get("ingested_at"),
        stored.get("created_at"),
    ):
        stored_dt = _parse_source_modified(raw)
        if stored_dt is not None:
            return (source_modified_dt - stored_dt).total_seconds() > 1.0
    return False


def _hashlib_sha1(value: str) -> str:
    import hashlib

    return hashlib.sha1(value.encode("utf-8")).hexdigest()


# Extraction budget: the cap on EXTRACTED TEXT per file (not on source
# records). The old hard caps (5 sheets x 100 rows, 50 PDF pages, 500 docx
# paragraphs, 1000 CSV rows) silently truncated real business files — live
# 2026-09-03, Consolidated Price List 2019.xlsx (46 sheets, ~2.7M chars)
# ingested without its machine-pricing sheets, so the agent could never find
# the WG350DSAV row and confabulated the price instead. The write path
# chunks arbitrarily long text (vector_upsert ::c{i} rows), so extraction
# may emit everything up to this budget; the budget only bounds memory and
# embedding cost for pathological files. Env-tunable per deployment.
# NOTE: the explicit read path (_read_storage_file) deliberately exceeds
# this — opening a named file must see all of it.
DEFAULT_EXTRACTION_MAX_CHARS = 4_000_000

# Safety ceiling for the explicit read path (a named-file open). Still
# finite so a hostile multi-GB workbook cannot OOM the process, but far
# above any real spreadsheet/PDF.
READ_EXTRACTION_MAX_CHARS = 50_000_000


def extraction_max_chars() -> int:
    try:
        return max(50_000, int(os.getenv("ATOM_EXTRACTION_MAX_CHARS", "")))
    except (TypeError, ValueError):
        return DEFAULT_EXTRACTION_MAX_CHARS


class _ExtractionBudget:
    """Accumulates text parts until the per-file char budget is exhausted.

    Callers append whole logical units (a sheet, a page, a row block) and
    check ``exhausted`` between units, so a unit is never split mid-way;
    ``truncation_note(total_units, consumed_units)`` renders the marker that
    tells recall (and the model) exactly what was skipped.
    """

    __slots__ = ("limit", "_parts", "_len")

    def __init__(self, limit: Optional[int] = None):
        self.limit = extraction_max_chars() if limit is None else limit
        self._parts: List[str] = []
        self._len = 0

    @property
    def consumed(self) -> int:
        return self._len

    @property
    def exhausted(self) -> bool:
        return self._len >= self.limit

    def add(self, part: str) -> bool:
        """Append if budget remains; True when written, False when dropped."""
        if self.exhausted or not part:
            return False
        self._parts.append(part)
        self._len += len(part)
        return True

    def join(self, sep: str = "\n") -> str:
        return sep.join(p for p in self._parts if p)

    def truncation_note(self, total: int, consumed: int) -> str:
        if consumed >= total:
            return ""
        return (
            f"... (extraction budget reached: showing {consumed} of {total} "
            f"sections; raise ATOM_EXTRACTION_MAX_CHARS to extract more)"
        )


class DocumentParser:
    """
    Parses various document formats and extracts text.
    Uses docling as primary parser with fallback to legacy parsers.
    Reuses existing parsers from DocumentLifecycleLearner where available.
    """
    
    # Docling processor (lazy-loaded)
    _docling_processor = None
    
    @classmethod
    def _get_docling_processor(cls):
        """Get or initialize the docling processor."""
        if cls._docling_processor is None:
            try:
                from core.docling_processor import get_docling_processor, is_docling_available
                if is_docling_available():
                    cls._docling_processor = get_docling_processor()
                    logger.info("Docling processor initialized for DocumentParser")
                else:
                    cls._docling_processor = False  # Mark as unavailable
            except ImportError:
                cls._docling_processor = False
                logger.debug("Docling not available, using fallback parsers")
        return cls._docling_processor if cls._docling_processor else None
    
    @staticmethod
    async def parse_document(file_content: bytes, file_type: str, file_name: str,
                             max_chars: Optional[int] = None) -> str:
        """Parse document and extract text content.

        ``max_chars`` overrides the per-file extraction budget for this call
        (None = the configured budget). The explicit read path passes a
        much larger ceiling: when the user opens a NAMED file, the answer
        must be able to see every sheet/page of it.
        """
        try:
            # Try docling first for supported formats
            docling = DocumentParser._get_docling_processor()
            docling_formats = ['pdf', 'docx', 'doc', 'pptx', 'ppt', 'xlsx', 'xls', 'html', 'htm', 'png', 'jpg', 'jpeg', 'tiff']

            if docling and file_type in docling_formats:
                try:
                    result = await docling.process_document(
                        source=file_content,
                        file_type=file_type,
                        file_name=file_name,
                        export_format="markdown"
                    )
                    if result.get("success") and result.get("content"):
                        logger.info(f"Docling parsed {file_name}: {result.get('total_chars', 0)} chars")
                        return result["content"]
                    else:
                        logger.warning(f"Docling parse failed for {file_name}, using fallback")
                except Exception as e:
                    logger.warning(f"Docling error for {file_name}: {e}, using fallback")

            # Fallback to legacy parsers
            if file_type in ["txt", "md", "toml", "yaml", "yml", "xml", "html", "ini", "cfg", "conf", "log", "sql", "py", "ts", "js", "sh", "bat", "env"]:
                return file_content.decode("utf-8", errors="ignore")

            elif file_type == "json":
                try:
                    data = json.loads(file_content.decode("utf-8", errors="ignore"))
                    return json.dumps(data, indent=2)
                except Exception:
                    # Unparseable JSON is not raw text worth embedding —
                    # returning the broken bytes verbatim pollutes memory
                    # with junk rows.
                    return ""

            elif file_type == "csv":
                return DocumentParser._parse_csv(file_content, max_chars=max_chars)

            elif file_type == "pdf":
                return await DocumentParser._parse_pdf(file_content, max_chars=max_chars)

            elif file_type in ["doc", "docx"]:
                return await DocumentParser._parse_docx(file_content, max_chars=max_chars)

            elif file_type in ["xlsx", "xls"]:
                return await DocumentParser._parse_excel(file_content, max_chars=max_chars)

            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return ""

        except Exception as e:
            logger.error(f"Failed to parse {file_name}: {e}")
            return ""
    
    @staticmethod
    def _parse_csv(content: bytes, file_path: str = None, workspace_id: str = "default", max_chars: Optional[int] = None) -> str:
        """Parse CSV to text - reuses DataIngestionService logic.
        Also extracts implicit formulas from column patterns.
        """
        # Formula memory from implicit column patterns (same extractor the
        # MCP upload path uses). Runs unconditionally — the extractor reads a
        # temp copy of the bytes, so gating on the caller's disk path only
        # ever excluded cloud-connector ingests (bytes-only).
        try:
            from core.formula_extractor import get_formula_extractor
            extractor = get_formula_extractor(workspace_id)
            # Need to save content to temp file
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".csv", delete=False, mode='wb') as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            try:
                extracted = extractor.extract_from_csv(tmp_path, auto_store=True)
                if extracted:
                    logger.info(f"Extracted {len(extracted)} formulas from CSV")
            finally:
                os.unlink(tmp_path)
        except Exception as fe:
            logger.warning(f"CSV formula extraction failed: {fe}")
        
        try:
            import csv
            text = content.decode("utf-8", errors="ignore")
            reader = csv.reader(io.StringIO(text))
            budget = _ExtractionBudget(limit=max_chars)
            total_rows = 0
            for row in reader:
                total_rows += 1
                if not budget.add(" | ".join(row)):
                    break
            note = budget.truncation_note(total_rows, len(budget._parts))
            return budget.join() + ("\n" + note if note else "")
        except Exception as e:
            logger.error(f"CSV parse error: {e}")
            return content.decode("utf-8", errors="ignore")

    
    @staticmethod
    async def _parse_pdf(content: bytes, max_chars: Optional[int] = None) -> str:
        """Parse PDF to text - compatible with DocumentLifecycleLearner"""
        try:
            # Use pypdf (PyPDF2 merged into pypdf package)
            import pypdf as PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            budget = _ExtractionBudget(limit=max_chars)
            total_pages = len(reader.pages)
            for page in reader.pages:
                if not budget.add(page.extract_text() or ""):
                    break
            note = budget.truncation_note(total_pages, len(budget._parts))
            return budget.join("\n\n") + ("\n\n" + note if note else "")
        except ImportError:
            logger.warning("pypdf not available, PDF parsing disabled")
            return "[PDF content - parser not available]"
        except Exception as e:
            logger.error(f"PDF parse error: {e}")
            return ""
    
    @staticmethod
    async def _parse_docx(content: bytes, max_chars: Optional[int] = None) -> str:
        """Parse DOCX to text - compatible with DocumentLifecycleLearner"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            budget = _ExtractionBudget(limit=max_chars)
            total_units = len(doc.paragraphs) + len(doc.tables)
            consumed = 0

            # Extract paragraphs
            for para in doc.paragraphs:
                if not budget.add(para.text):
                    break
                consumed += 1

            # Also extract tables (from DocumentLifecycleLearner) — table text
            # shares the budget with paragraphs, so a paragraph-heavy doc
            # leaves less room for tables and vice versa.
            if not budget.exhausted:
                for table in doc.tables:
                    done = False
                    for row in table.rows:
                        if not budget.add(" | ".join(cell.text for cell in row.cells)):
                            done = True
                            break
                    consumed += 1
                    if done or budget.exhausted:
                        break

            note = budget.truncation_note(total_units, consumed)
            return budget.join() + ("\n" + note if note else "")
        except ImportError:
            logger.warning("python-docx not available")
            return "[DOCX content - parser not available]"
        except Exception as e:
            logger.error(f"DOCX parse error: {e}")
            return ""
    
    @staticmethod
    async def _parse_excel(content: bytes, file_path: str = None, workspace_id: str = "default", max_chars: Optional[int] = None) -> str:
        """Parse Excel to text - compatible with DocumentLifecycleLearner.
        Also extracts formulas and stores them in Atom's formula memory.
        """
        # Old-format .xls (OLE2 Compound File): only xlrd reads these. The
        # zip-based parsers (openpyxl/pandas) raise BadZipFile on OLE2 content,
        # so detect by the OLE2 magic bytes BEFORE anything touches them.
        if content[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
            try:
                import xlrd
                wb = xlrd.open_workbook(file_contents=content)
                budget = _ExtractionBudget(limit=max_chars)
                sheets_done = 0
                summaries: List[str] = []
                for sheet in wb.sheets():
                    if budget.exhausted:
                        break
                    budget.add(f"--- Sheet: {sheet.name} ---")
                    ncols = getattr(sheet, "ncols", 0)
                    if ncols:
                        headers = []
                        for c in range(min(ncols, 40)):
                            h = " ".join(str(sheet.cell_value(0, c)).split())[:48]
                            if h:
                                headers.append(f"{chr(65 + c) if c < 26 else '?'}={h}")
                        if headers:
                            budget.add("COLS: " + " | ".join(headers))
                    for r in range(1, getattr(sheet, "nrows", 0)):
                        row = []
                        for c in range(ncols):
                            val = sheet.cell_value(r, c)
                            row.append(str(val) if val is not None else "")
                        if not budget.add(f"R{r + 1} | " + " | ".join(row)):
                            break
                    sheets_done += 1
                    summaries.append(f"{sheet.name}: {getattr(sheet, 'nrows', 0)} rows, {ncols} cols")
                note = budget.truncation_note(len(wb.sheets()), sheets_done)
                index = "\n".join([f"WORKBOOK INDEX: {len(summaries)} sheets"] + [f"- {s}" for s in summaries])
                return index + "\n" + budget.join() + ("\n" + note if note else "")
            except ImportError:
                logger.warning("xlrd not installed; cannot parse old .xls (OLE2) files")
                return ""
            except Exception as e:
                logger.error(f"XLS (xlrd) parse error: {e}")
                return ""

        # Formula memory (Phase 19 mechanism — same extractor the MCP upload
        # path uses). Runs unconditionally: the extractor reads a temp copy
        # of the bytes, so there is no reason to gate it on the caller having
        # a disk path. Cloud-connector ingests (bytes-only) were silently
        # excluded here, so formula memory never saw any WorkDrive/OneDrive/
        # Drive workbook. Hostile workbooks fail fast inside the extractor
        # (openpyxl) and are fault-isolated to [].
        try:
            from core.formula_extractor import get_formula_extractor
            extractor = get_formula_extractor(workspace_id)
            # openpyxl needs a real file, not bytes
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            try:
                extracted = extractor.extract_from_excel(tmp_path, auto_store=True)
                if extracted:
                    logger.info(f"Extracted {len(extracted)} formulas from Excel")
            finally:
                os.unlink(tmp_path)
        except Exception as fe:
            logger.warning(f"Formula extraction failed: {fe}")
        
        try:
            import pandas as pd

            # Read ALL sheets — sheet-count caps silently dropped real pricing
            # sections (live 2026-09-03: Consolidated Price List 2019.xlsx
            # ingested without its machine sheets; only the first five sheets
            # were extracted). Only total extracted chars are bounded.
            xls = pd.ExcelFile(io.BytesIO(content))
            budget = _ExtractionBudget(limit=max_chars)
            sheets_done = 0
            summaries: List[str] = []
            for sheet_name in xls.sheet_names:
                if budget.exhausted:
                    break
                df = pd.read_excel(xls, sheet_name=sheet_name)
                budget.add(f"--- Sheet: {sheet_name} ---")
                # Structure anchoring: column letters -> headers, and the
                # sheet-row offset (pandas index 0 == sheet row 2).
                headers = []
                for i, col in enumerate(df.columns[:40]):
                    h = " ".join(str(col).split())[:48]
                    if h and h.lower() != "nan":
                        letters = (
                            chr(65 + i) if i < 26 else f"A{chr(65 + i - 26)}"
                        )
                        headers.append(f"{letters}={h}")
                if headers:
                    budget.add("COLS: " + " | ".join(headers))
                if len(df):
                    budget.add(f"ROWS: sheet rows 2..{len(df) + 1}")
                # Rows in sheet coordinates (pandas index 0 == sheet row 2),
                # same R# convention as the raw-XML path so a citation is
                # unambiguous about WHERE the value lives. Values stay
                # positional — the COLS line above is the single schema
                # anchor (header=value per row would re-serialize the header
                # on every row: the token cost SheetCompressor exists to
                # avoid).
                for idx, row_vals in enumerate(df.itertuples(index=False, name=None)):
                    cells = ["" if v is None else " ".join(str(v).split()) for v in row_vals]
                    while cells and not cells[-1]:
                        cells.pop()
                    if not budget.add(f"R{idx + 2} | " + " | ".join(cells)):
                        break
                sheets_done += 1
                summaries.append(
                    f"{sheet_name}: {len(df)} data rows, {len(df.columns)} cols"
                    + (f" | headers: {', '.join(str(c) for c in df.columns[:8])}" if len(df.columns) else "")
                )
            note = budget.truncation_note(len(xls.sheet_names), sheets_done)
            index = "\n".join([f"WORKBOOK INDEX: {len(summaries)} sheets"] + [f"- {s}" for s in summaries])
            return index + "\n" + budget.join() + ("\n" + note if note else "")
        except ImportError:
            # Fallback to openpyxl
            try:
                from openpyxl import load_workbook
                wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
                budget = _ExtractionBudget(limit=max_chars)
                sheets_done = 0
                summaries: List[str] = []
                for sheet_name in wb.sheetnames:
                    if budget.exhausted:
                        break
                    sheet = wb[sheet_name]
                    budget.add(f"=== Sheet: {sheet_name} ===")
                    rownum = 0
                    headers = []
                    for row in sheet.iter_rows(values_only=True):
                        rownum += 1
                        cells = [str(cell) if cell is not None else "" for cell in row]
                        if rownum == 1:
                            for i, h in enumerate(cells[:40]):
                                h = " ".join(h.split())[:48]
                                if h:
                                    letters = (
                                        chr(65 + i) if i < 26 else f"A{chr(65 + i - 26)}"
                                    )
                                    headers.append(f"{letters}={h}")
                            if headers:
                                budget.add("COLS: " + " | ".join(headers))
                            continue
                        if not budget.add(f"R{rownum} | " + " | ".join(cells)):
                            break
                    sheets_done += 1
                    summaries.append(f"{sheet_name}: {rownum} rows" + (f" | headers: {', '.join(h.split('=',1)[-1] for h in headers[:8])}" if headers else ""))
                note = budget.truncation_note(len(wb.sheetnames), sheets_done)
                index = "\n".join([f"WORKBOOK INDEX: {len(summaries)} sheets"] + [f"- {s}" for s in summaries])
                return index + "\n" + budget.join() + ("\n" + note if note else "")
            except ImportError:
                logger.warning("No Excel parser available")
                return "[Excel content - parser not available]"
        except Exception as e:
            # openpyxl is strict about workbook XML schema and hard-fails on
            # spreadsheets written by non-Excel tools (Zoho Sheet exports trip
            # read_strings on unusual sharedStrings content). Those files are
            # still plain well-formed XML — extract them directly instead of
            # returning "" (which the caller turns into a no_text skip and the
            # file silently never becomes searchable).
            logger.error(f"Excel parse error: {e}")
            try:
                fallback_text = DocumentParser._parse_xlsx_raw(content, max_chars=max_chars)
                if fallback_text and fallback_text.strip():
                    logger.info(
                        f"Raw-XML fallback extracted {len(fallback_text)} chars "
                        f"after Excel parser failure"
                    )
                    return fallback_text
            except Exception as raw_err:
                logger.error(f"Raw-XML Excel fallback failed: {raw_err}")
            return ""

    @staticmethod
    def _parse_xlsx_raw(content: bytes, max_sheets: Optional[int] = None, max_rows: Optional[int] = None, max_chars: Optional[int] = None) -> str:
        """Extract cell text from an xlsx zip with stdlib XML parsing only.

        Engine-agnostic last resort for workbooks openpyxl/pandas reject:
        reads sharedStrings.xml (concatenating rich-text runs, ignoring
        structure it doesn't understand), inline strings, and raw values.
        Namespace-agnostic so producer-specific namespaces don't matter.
        All sheets/rows are extracted up to the shared per-file char budget;
        ``max_sheets``/``max_rows`` remain as optional additional caps for
        callers that want them.

        Structure preservation (spreadsheet-RAG practice — SpreadsheetLLM/
        SheetCompressor arXiv:2407.09025 structure anchoring; TableRAG
        arXiv:2410.04739 schema-first retrieval):
          - a WORKBOOK INDEX at the head: per sheet — name, row/col counts,
            header names, formula count ("which sheet has prices?" is
            answerable from one chunk);
          - a per-sheet column map (column letter -> header) — the schema
            anchor that makes rows interpretable without serializing every
            cell address;
          - real row numbers on every row (``R17 | ...``) so any hit cites
            its sheet row ("LINMAC R17") instead of an anonymous tuple;
          - per-cell formulas kept next to their cached values.
        """
        import zipfile
        import xml.etree.ElementTree as ET

        def _local(tag: str) -> str:
            return tag.rsplit("}", 1)[-1]

        def _texts(si) -> str:
            # <si> children: <t> plain text, <r> rich-text runs, <rPh>
            # phonetic hints (pronunciations — not content, skip them).
            parts = []
            for child in si:
                tag = _local(child.tag)
                if tag == "t":
                    parts.append(child.text or "")
                elif tag == "r":
                    parts.extend(
                        sub.text or "" for sub in child if _local(sub.tag) == "t"
                    )
            return "".join(parts)

        with zipfile.ZipFile(io.BytesIO(content)) as zf:
            shared: list = []
            if "xl/sharedStrings.xml" in zf.namelist():
                sst_root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
                for si in sst_root:
                    if _local(si.tag) == "si":
                        shared.append(_texts(si))

            # Real sheet names from workbook.xml + its rels (rid -> part
            # target). "Sheet: sheet1.xml" is retrieval noise — the name is
            # often the only semantic label a price-book tab has.
            names: Dict[str, str] = {}
            try:
                wb_root = ET.fromstring(zf.read("xl/workbook.xml"))
                rels_root = ET.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
                rid_to_target = {
                    rel.get("Id"): rel.get("Target", "")
                    for rel in rels_root
                    if _local(rel.tag) == "Relationship"
                }
                for sheet in wb_root.iter():
                    if _local(sheet.tag) != "sheet":
                        continue
                    rid = sheet.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                    )
                    target = rid_to_target.get(rid, "")
                    target = target.lstrip("/")
                    if target and not target.startswith("xl/"):
                        target = f"xl/{target}"
                    if target:
                        names[target] = sheet.get("name") or target
            except Exception as name_err:  # noqa: BLE001 — names are best-effort
                logger.debug(f"workbook sheet-name read failed: {name_err}")

            budget = _ExtractionBudget(limit=max_chars)
            sheets = sorted(n for n in zf.namelist() if n.startswith("xl/worksheets/") and n.endswith(".xml"))

            def _col_letters(ref: str) -> str:
                """'C17' -> 'C' — the column letters of a cell reference."""
                return "".join(ch for ch in (ref or "") if ch.isalpha())

            sheet_summaries: List[str] = []
            sheets_done = 0
            for sheet_path in sheets:
                if max_sheets is not None and sheets_done >= max_sheets:
                    break
                if budget.exhausted:
                    break
                sheet_name = names.get(sheet_path, sheet_path.rsplit('/', 1)[-1])
                budget.add(f"=== Sheet: {sheet_name} ===")
                root = ET.fromstring(zf.read(sheet_path))
                rows_shown = 0
                row_numbers_seen = 0
                formula_count = 0
                header_map: List[str] = []
                for row in root.iter():
                    if _local(row.tag) != "row":
                        continue
                    if max_rows is not None and rows_shown >= max_rows:
                        budget.add("... (truncated)")
                        break
                    cells = []
                    row_has_formula = False
                    for c in row:
                        if _local(c.tag) != "c":
                            continue
                        ctype = c.get("t")
                        if ctype == "inlineStr":
                            cells.append(
                                "".join(
                                    node.text or ""
                                    for node in c.iter()
                                    if _local(node.tag) == "t"
                                )
                            )
                            continue
                        v = next((ch for ch in c if _local(ch.tag) == "v"), None)
                        if v is None:
                            cells.append("")
                        elif ctype == "s":
                            try:
                                cells.append(shared[int(v.text)])
                            except (ValueError, IndexError):
                                cells.append(v.text or "")
                        else:
                            cells.append(v.text or "")
                        # Formula cells: <v> holds only the last computed
                        # value; the <f> element is the actual business logic
                        # (markups, currency conversion). Render both — the
                        # cached value alone hides HOW the number is derived.
                        f_el = next(
                            (ch for ch in c if _local(ch.tag) == "f"), None
                        )
                        if f_el is not None and (f_el.text or "").strip():
                            cells[-1] = f"{cells[-1]} [={f_el.text.strip()}]".strip()
                            row_has_formula = True
                    while cells and not str(cells[-1]).strip():
                        cells.pop()  # trailing empties are noise
                    # Structure anchoring: the first substantive row of the
                    # sheet defines the column map (letter -> header).
                    if not header_map and sum(1 for c in cells if str(c).strip()) >= 2:
                        for c in row:
                            if _local(c.tag) != "c":
                                continue
                            if len(header_map) >= 40:
                                break
                            v = next((ch for ch in c if _local(ch.tag) == "v"), None)
                            header = ""
                            if c.get("t") == "s" and v is not None and v.text:
                                try:
                                    header = shared[int(v.text)]
                                except (ValueError, IndexError):
                                    header = ""
                            elif c.get("t") == "inlineStr":
                                header = "".join(
                                    node.text or "" for node in c.iter()
                                    if _local(node.tag) == "t"
                                )
                            header = " ".join(str(header).split())[:48]
                            if header:
                                header_map.append(f"{_col_letters(c.get('r') or '')}={header}")
                    rownum = row.get("r") or str(rows_shown + 1)
                    row_numbers_seen += 1
                    if row_has_formula:
                        formula_count += 1
                    prefix = f"R{rownum} | "
                    if not budget.add(prefix + " | ".join(str(c) for c in cells)):
                        break
                    rows_shown += 1
                if header_map:
                    budget.add("COLS: " + " | ".join(header_map))
                sheets_done += 1
                sheet_summaries.append(
                    f"{sheet_name}: {row_numbers_seen} rows, "
                    f"{formula_count} formula rows"
                    + (f" | headers: {', '.join(h.split('=', 1)[-1] for h in header_map[:8])}" if header_map else "")
                )
            # Workbook index (TableRAG-style schema-first discovery): one
            # chunk at the head that answers "which sheet has X" without
            # walking 4M chars.
            index_lines = [f"WORKBOOK INDEX: {len(sheet_summaries)} sheets"]
            index_lines.extend(f"- {s}" for s in sheet_summaries)
            text = budget.join() + ("\n" + budget.truncation_note(len(sheets), sheets_done) if sheets_done < len(sheets) else "")
            text = "\n".join(index_lines) + "\n" + text
            return text


class AutoDocumentIngestionService:
    """
    Manages automatic document ingestion from connected file storage integrations.
    
    Features:
    - Per-integration settings in user preferences
    - Auto-sync new/updated files
    - Parse multiple file formats
    - Ingest to Atom Memory (LanceDB + GraphRAG)
    """
    
    def __init__(self, workspace_id: str = "default"):
        # Per-workspace construction: settings, ingested-doc cache, LanceDB
        # handler and durable settings rows are all workspace-scoped. The old
        # fixed-"default" singleton meant every non-default workspace's sync
        # read and wrote the default workspace's stores.
        self.workspace_id = (workspace_id or "default").strip() or "default"
        self.settings: Dict[str, IngestionSettings] = {}
        self.ingested_docs: Dict[str, IngestedDocument] = {}  # key = external_id
        self.parser = DocumentParser()
        self._running = False

        # Initialize memory handler
        try:
            from core.lancedb_handler import get_lancedb_handler
            self.memory_handler = get_lancedb_handler(self.workspace_id)
        except ImportError:
            self.memory_handler = None
            logger.warning("LanceDB handler not available")
        
        # Initialize secrets redactor
        try:
            from core.secrets_redactor import get_secrets_redactor
            self.redactor = get_secrets_redactor()
        except ImportError:
            self.redactor = None
    
    def get_settings(self, integration_id: str) -> IngestionSettings:
        """Get or create settings for an integration"""
        if integration_id not in self.settings:
            settings = IngestionSettings(
                integration_id=integration_id,
                workspace_id=self.workspace_id
            )
            # Hydrate from the durable row if one exists (restart survival).
            self._load_settings_row(integration_id, settings)
            self.settings[integration_id] = settings
        return self.settings[integration_id]

    def _settings_row_query(self, db, integration_id: str):
        from core.models import IngestionSettings as IngestionSettingsRow

        return (
            db.query(IngestionSettingsRow)
            .filter(
                IngestionSettingsRow.workspace_id == self.workspace_id,
                IngestionSettingsRow.integration_id == integration_id,
            )
            .first()
        )

    @staticmethod
    def _settings_persistence_enabled() -> bool:
        """Same kill switch as hybrid ingestion state (tests disable it)."""
        return os.getenv("ATOM_INGESTION_PERSIST_STATE", "true").lower() in (
            "1", "true", "yes",
        )

    def _load_settings_row(self, integration_id: str, settings: IngestionSettings) -> None:
        """Best-effort hydration from the ``ingestion_settings`` table.

        Settings previously lived only in this process's memory — a restart
        wiped enabled flags, folders and last_sync. Never raises."""
        if not self._settings_persistence_enabled():
            return
        try:
            from core.database import get_db_session

            with get_db_session() as db:
                row = self._settings_row_query(db, integration_id)
                if row is None:
                    return
                settings.enabled = bool(row.enabled)
                settings.auto_sync_new_files = bool(row.auto_sync_new_files)
                settings.file_types = list(row.file_types or settings.file_types)
                settings.sync_folders = list(row.sync_folders or [])
                settings.exclude_folders = list(row.exclude_folders or [])
                settings.max_file_size_mb = row.max_file_size_mb or settings.max_file_size_mb
                settings.sync_frequency_minutes = (
                    row.sync_frequency_minutes or settings.sync_frequency_minutes
                )
                if row.last_sync is not None:
                    # SQLite returns naive datetimes; sync math compares
                    # against timezone-aware now().
                    settings.last_sync = (
                        row.last_sync.replace(tzinfo=timezone.utc)
                        if row.last_sync.tzinfo is None
                        else row.last_sync
                    )
        except Exception as e:
            logger.debug(f"Ingestion settings load skipped for {integration_id}: {e}")

    def _persist_settings(self, settings: IngestionSettings) -> None:
        """Best-effort durable copy of ingestion settings. Never raises.

        Only the document-ingestion columns are touched — hybrid ingestion
        stores its pipeline state (entity_types, sync_mode, …) in the same
        rows and must survive this upsert."""
        if not self._settings_persistence_enabled():
            return
        try:
            from core.database import get_db_session

            with get_db_session() as db:
                row = self._settings_row_query(db, settings.integration_id)
                if row is None:
                    from core.models import IngestionSettings as IngestionSettingsRow

                    row = IngestionSettingsRow(
                        workspace_id=self.workspace_id,
                        integration_id=settings.integration_id,
                    )
                    db.add(row)
                row.enabled = settings.enabled
                row.auto_sync_new_files = settings.auto_sync_new_files
                row.file_types = list(settings.file_types or [])
                row.sync_folders = list(settings.sync_folders or [])
                row.exclude_folders = list(settings.exclude_folders or [])
                row.max_file_size_mb = settings.max_file_size_mb
                row.sync_frequency_minutes = settings.sync_frequency_minutes
                row.last_sync = settings.last_sync
                db.commit()
        except Exception as e:
            logger.warning(
                f"Failed to persist ingestion settings for {settings.integration_id}: {e}"
            )
    
    def update_settings(
        self,
        integration_id: str,
        enabled: Optional[bool] = None,
        auto_sync_new_files: Optional[bool] = None,
        file_types: Optional[List[str]] = None,
        sync_folders: Optional[List[str]] = None,
        exclude_folders: Optional[List[str]] = None,
        max_file_size_mb: Optional[int] = None,
        sync_frequency_minutes: Optional[int] = None,
    ) -> IngestionSettings:
        """Update settings for an integration"""
        settings = self.get_settings(integration_id)
        
        if enabled is not None:
            settings.enabled = enabled
        if auto_sync_new_files is not None:
            settings.auto_sync_new_files = auto_sync_new_files
        if file_types is not None:
            settings.file_types = file_types
        if sync_folders is not None:
            settings.sync_folders = sync_folders
        if exclude_folders is not None:
            settings.exclude_folders = exclude_folders
        if max_file_size_mb is not None:
            settings.max_file_size_mb = max_file_size_mb
        if sync_frequency_minutes is not None:
            settings.sync_frequency_minutes = sync_frequency_minutes
        
        logger.info(f"Updated ingestion settings for {integration_id}: enabled={settings.enabled}")
        self._persist_settings(settings)
        return settings

    async def process_file_bytes(
        self,
        content: bytes,
        file_name: str,
        source: str = "upload",
        user_id: str = "system",
        workspace_id: Optional[str] = None,
        role: Optional[str] = None,
        extra_metadata: Optional[Dict[str, Any]] = None,
        external_id: Optional[str] = None,
        explicit: bool = True,
    ) -> Dict[str, Any]:
        """Parse raw file bytes and ingest the extracted text into Atom memory.

        Shared ingestion path used by cloud-drive connectors (OneDrive, Zoho
        WorkDrive) and direct uploads. Parses with DocumentParser, redacts
        secrets, and writes to LanceDB (documents table) with knowledge
        extraction enabled.

        Args:
            content: Raw file bytes.
            file_name: File name (used to infer type and metadata).
            source: Source label (e.g. "onedrive", "zoho_workdrive").
            user_id: Owning user id.
            workspace_id: Optional workspace override (defaults to service default).
            role: Optional AI-employee role (AgentRegistry.category, lowercased)
                the ingested file is relevant to. Tagged in the LanceDB metadata
                so role-aware recall (WorldModelService) surfaces it to the right
                employee's memory. None/empty = general knowledge.
            external_id: SOURCE-NATIVE unique id (Drive fileId, OneDrive
                driveItem id, Box file id, Dropbox path …). Preferred identity —
                titles are not identity. Falls back to a SHA-256 of the
                extracted text (content addressing) when absent.
            explicit: False for AUTOMATIC bulk syncs — lets the integration's
                content mode (hybrid/list_only) skip content ingestion to save
                disk + extraction cost. Explicit user/agent pulls always pass
                True and are never mode-gated.

        Returns:
            Dict with ``status``, ``file_name``, ``chars_ingested``, ``doc_id``.
        """
        file_ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
        if not file_ext:
            return {"status": "skipped", "reason": "no_file_extension", "file_name": file_name}

        _external_id = str(
            external_id or (extra_metadata or {}).get("external_id") or ""
        ).strip()
        # Key normalization: connectors name the modified time inconsistently
        # (WorkDrive's bulk walker passes "modified_at"; uploads/tools pass
        # "source_modified_at"). This timestamp IS the update-detection
        # signal — reading only one key meant WorkDrive files never stamped
        # it, so hybrid-mode refresh probes ("has the source changed since we
        # ingested?") always answered False and source-side edits were
        # silently skipped in hybrid/list_only mode.
        source_modified_dt = _parse_source_modified(
            (extra_metadata or {}).get("source_modified_at")
            or (extra_metadata or {}).get("modified_at")
        )

        # Content-mode gate for storage drives: hybrid/list_only keep the
        # file index but skip automatic content ingestion. Explicit
        # user/agent pulls bypass this — that is the entire point of hybrid.
        # EXCEPTION — source-side updates: when the walker reports a newer
        # modified_at for a file whose content we ALREADY hold, the stored
        # copy is silently rotting (it still reads "fresh" to recall while
        # the source moved on). Refresh it: the walker already downloaded
        # these bytes, so the marginal cost is one parse+write, and the
        # upsert replaces the old row under the same identity key.
        stored_is_outdated = False
        if not explicit and source and source != "upload" and _external_id:
            try:
                probe_id = f"ext_{_hashlib_sha1(f'{source}:{_external_id}')[:24]}"
                stored = await asyncio.to_thread(
                    self.memory_handler.get_document_by_id, "documents", probe_id
                )
                if stored is None:
                    # chunked layout: the family lives under {doc_id}::c{i}
                    stored = await asyncio.to_thread(
                        self.memory_handler.get_document_by_id,
                        "documents",
                        f"{probe_id}::c0",
                    )
                stored_is_outdated = _stored_copy_older_than(
                    stored, source_modified_dt
                )
            except Exception as probe_err:  # noqa: BLE001 — probe is best-effort
                logger.debug(f"update probe skipped for {file_name}: {probe_err}")

        if not explicit and source and source != "upload":
            try:
                from core.hybrid_data_ingestion import get_hybrid_ingestion_service

                mode = get_hybrid_ingestion_service(
                    workspace_id or "default"
                ).get_content_mode(source)
                if mode in ("hybrid", "list_only") and not stored_is_outdated:
                    logger.info(
                        f"Content-mode {mode}: skipping auto-ingest of {file_name} "
                        f"from {source} (metadata indexed; content on demand)"
                    )
                    return {
                        "status": "skipped",
                        "reason": f"content_mode_{mode}",
                        "file_name": file_name,
                    }
                if mode in ("hybrid", "list_only") and stored_is_outdated:
                    logger.info(
                        f"Content refresh: {file_name} changed at source — "
                        f"re-ingesting stored copy (was outdated)"
                    )
            except Exception as mode_err:
                logger.debug(f"Content-mode lookup failed for {source}: {mode_err}")

        try:
            text = await self.parser.parse_document(content, file_ext, file_name)
        except Exception as parse_err:
            logger.warning(f"Failed to parse {file_name} ({file_ext}): {parse_err}")
            return {"status": "error", "reason": "parse_failed", "file_name": file_name}

        # Blank-only text is junk; short-but-real content (e.g. "data") is a
        # valid document and must not be dropped.
        if not text or not text.strip():
            return {"status": "skipped", "reason": "no_text", "file_name": file_name}

        # Redact secrets before storage
        if self.redactor:
            try:
                redaction = self.redactor.redact(text)
                if getattr(redaction, "has_secrets", False):
                    logger.info(f"Redacted secrets from {file_name}")
                    text = redaction.redacted_text
            except Exception as redact_err:
                logger.debug(f"Secrets redaction skipped for {file_name}: {redact_err}")

        ws_id = workspace_id or self.workspace_id
        chars_ingested = 0

        # Per-workspace handler: the workspace override previously only
        # stamped metadata — the row still landed in the default workspace's
        # store, invisible to that workspace's recall.
        _handler = self.memory_handler
        if ws_id and ws_id != self.workspace_id:
            try:
                if not hasattr(self, "_ws_handlers"):
                    self._ws_handlers: Dict[str, Any] = {}
                if ws_id not in self._ws_handlers:
                    from core.lancedb_handler import get_lancedb_handler

                    self._ws_handlers[ws_id] = get_lancedb_handler(ws_id)
                _handler = self._ws_handlers[ws_id]
            except Exception as ws_handler_err:  # noqa: BLE001 — fall back to default
                logger.warning(
                    f"workspace handler unavailable for {ws_id}, using default: {ws_handler_err}"
                )
                _handler = self.memory_handler

        # Join-key bridge (hybrid search, Step 1): the file-ingest path creates
        # no PG IngestedDocument row, so vector hits from here can't resolve to
        # documents.cat. Stamp a stable doc_id + source_type:"file" so the
        # hybrid service can flag these as bridged:false (no PG row) rather
        # than silently returning unresolvable hits.
        # Identity key (idempotency contract): prefer the SOURCE-NATIVE id the
        # connector passes (external_id param, else extra_metadata fallback);
        # otherwise fall back to a SHA-256 of the extracted text
        # (content-addressing: identical content = one row, any filename).
        # The id is SOURCE-SCOPED — two integrations may reuse the same raw
        # external id string, and one file's refresh must never delete the
        # other's row. Never key on file NAME.
        import hashlib as _hashlib

        from core.doc_freshness_service import (
            extraction_content_hash,
            extra_columns_for_ingest,
            has_current_extraction_version,
        )

        _content_hash = extraction_content_hash(text)
        if _external_id:
            _identity_input = f"{source}:{_external_id}"
            _file_doc_id = f"ext_{_hashlib.sha1(_identity_input.encode('utf-8')).hexdigest()[:24]}"
        else:
            _file_doc_id = "doc_" + _hashlib.sha256(text.encode("utf-8")).hexdigest()[:24]

        _meta: Dict[str, Any] = {
            "file_name": file_name,
            "file_type": file_ext,
            "file_size": len(content),
            "integration_id": source,
            "ingested_at": datetime.now(timezone.utc).isoformat(),
            "pg_document_id": _file_doc_id,
            "source_type": "file",
            "source_content_hash": _content_hash,
            "freshness_status": "fresh",
        }
        # Normalized source modified time — the comparison key for future
        # update detection (see the content-refresh exception at the gate).
        if source_modified_dt:
            _meta["source_modified_at"] = source_modified_dt.isoformat()
        if _external_id:
            _meta["external_id"] = _external_id
        # Connector-supplied context (e.g. WorkDrive folder path / root)
        if extra_metadata:
            _meta.update({k: v for k, v in extra_metadata.items() if v})
        # AI-employee relevance tag (Round 80): lets role-aware recall surface
        # this file to the employee whose role it was ingested for.
        if role:
            _meta["role"] = str(role).lower()

        if _handler:
            try:
                # Shared upsert contract (hash-skip / delete prior / write),
                # chunked: long documents store as {doc_id}::c{i} rows so
                # each region gets its own embedding — a single 55k-char row
                # embedded once matches only whatever its head looked like.
                # Short texts delegate to the plain single-row upsert.
                from core.vector_upsert import upsert_document_chunks

                _upsert_status = await upsert_document_chunks(
                    _handler,
                    table_name="documents",
                    text=text,
                    doc_id=_file_doc_id,
                    source=f"{source}:{file_name}",
                    metadata=_meta,
                    user_id=user_id,
                    workspace_id=ws_id,
                    extra_columns=extra_columns_for_ingest(
                        freshness_status="fresh",
                        source_modified_at=source_modified_dt,
                        source_url=None,
                    ),
                )
                if _upsert_status == "written":
                    chars_ingested = len(text)
                    logger.info(f"Ingested {file_name} ({chars_ingested} chars) from {source}")
                    # Aligned PG mirror row: id == the vector doc_id (join-key
                    # bridge). Without it this path's rows are invisible to the
                    # Knowledge VFS ls/grep and the lexical search leg — the
                    # Aug 2026 journey trace found 69 connector ingests with 0
                    # PG rows. Best-effort: never fail the ingest for the mirror.
                    try:
                        self._mirror_pg_row(
                            doc_id=_file_doc_id,
                            workspace_id=ws_id,
                            file_name=file_name,
                            file_path=f"{source}:{file_name}",
                            file_type=file_ext,
                            integration_id=source,
                            file_size_bytes=len(content),
                            content_preview=text[:500],
                            external_id=_external_id or f"vector:{_file_doc_id}",
                            content_hash=_content_hash,
                            role=str(_meta.get("role")) if _meta.get("role") else None,
                            source_modified_at=source_modified_dt,
                        )
                    except Exception as mirror_err:  # noqa: BLE001 — mirror is best-effort
                        logger.warning(f"PG mirror row skipped for {_file_doc_id}: {mirror_err}")
                else:
                    # Distinguish "we already have this exact content" from
                    # "the write FAILED" — both previously surfaced as
                    # 'unchanged', which told the user (and the agent) the
                    # file was already stored when the store had actually
                    # rejected the write.
                    reason = (
                        "unchanged" if _upsert_status == "skipped_unchanged"
                        else f"write_failed ({_upsert_status})"
                    )
                    return {
                        "status": "skipped",
                        "reason": reason,
                        "file_name": file_name,
                        "chars_ingested": 0,
                        "source": source,
                        "doc_id": _file_doc_id,
                    }
            except Exception as ingest_err:
                logger.error(f"Failed to ingest {file_name} to memory: {ingest_err}")
                return {"status": "error", "reason": "ingest_failed", "file_name": file_name}
        else:
            logger.warning("No memory_handler available; file content not stored")

        return {
            "status": "ingested" if chars_ingested else "skipped",
            "file_name": file_name,
            "chars_ingested": chars_ingested,
            "source": source,
            "doc_id": _file_doc_id,
        }

    async def sync_integration(
        self, 
        integration_id: str,
        force: bool = False
    ) -> Dict[str, Any]:
        """
        Sync documents from an integration.
        
        Returns:
            Dict with sync results
        """
        from core.doc_freshness_service import has_current_extraction_version

        settings = self.get_settings(integration_id)
        
        if not settings.enabled and not force:
            return {"skipped": True, "reason": "Integration not enabled"}
        
        # Check if sync is due
        if not force and settings.last_sync:
            minutes_since = (datetime.now(timezone.utc) - settings.last_sync).total_seconds() / 60
            if minutes_since < settings.sync_frequency_minutes:
                return {"skipped": True, "reason": "Recently synced"}
        
        logger.info(f"Starting document sync for {integration_id}")
        
        results = {
            "integration_id": integration_id,
            "started_at": datetime.now(timezone.utc).isoformat(),
            "files_found": 0,
            "files_ingested": 0,
            "files_skipped": 0,
            "errors": [],
            "newly_ingested_files": []
        }
        
        try:
            # Fetch file list from integration
            files = await self._list_files(integration_id, settings)
            results["files_found"] = len(files)

            # Track which external_ids the source currently reports. Used after
            # the loop to tombstone docs that were deleted upstream.
            seen_external_ids: Set[str] = set()

            for file_info in files:
                # LAMBDA SAFEGUARD: Check if we are approaching timeout (10 mins)
                # If running longer than 10 minutes, stop and let the next scheduled run pick up the rest
                if (datetime.now(timezone.utc) - datetime.fromisoformat(results["started_at"])).total_seconds() > 600:
                    logger.warning(f"Ingestion time limit reached (10m) for {integration_id}. Stopping early.")
                    results["errors"].append("Time limit reached - continuing in next run")
                    break

                try:
                    # Skip if already ingested and not modified AND the stored
                    # extraction was produced by the current extractor. The
                    # second condition is what makes extractor improvements
                    # self-propagating: after EXTRACTION_VERSION bumps, the
                    # next sync re-downloads (once) and re-extracts files the
                    # old extractor truncated, instead of skipping them forever
                    # because the SOURCE never changed (live 2026-09-03:
                    # Consolidated Price List 2019.xlsx stored without its
                    # machine-pricing sheets — the 5-sheet cap — and no sync
                    # would ever revisit it).
                    external_id = file_info.get("id")
                    if external_id:
                        seen_external_ids.add(external_id)
                    existing: Optional[IngestedDocument] = None
                    if external_id in self.ingested_docs:
                        existing = self.ingested_docs[external_id]
                        if (
                            file_info.get("modified_at") == existing.external_modified_at
                            and has_current_extraction_version(existing.source_content_hash)
                        ):
                            results["files_skipped"] += 1
                            continue
                        # Source modified_at differs → the stored copy is stale
                        # even if the re-download below later fails. Record it
                        # now so retrieval can downrank it.
                        try:
                            self._mark_doc_stale(existing, reason="source_modified_at_changed")
                        except Exception as stale_err:
                            logger.warning(f"Failed to mark stale for {external_id}: {stale_err}")
                    
                    # Check file type
                    file_ext = file_info.get("name", "").split(".")[-1].lower()
                    if file_ext not in settings.file_types:
                        results["files_skipped"] += 1
                        continue
                    
                    # Check file size
                    file_size = file_info.get("size", 0)
                    if file_size > settings.max_file_size_mb * 1024 * 1024:
                        results["files_skipped"] += 1
                        continue
                    
                    # Download and parse
                    content = await self._download_file(integration_id, file_info)
                    if not content:
                        continue
                    
                    text = await self.parser.parse_document(content, file_ext, file_info.get("name"))
                    if not text:
                        continue
                    
                    # Redact secrets before storage
                    if self.redactor:
                        redaction = self.redactor.redact(text)
                        if redaction.has_secrets:
                            logger.warning(f"Redacted {len(redaction.redactions)} secrets from {file_info.get('name')}")
                            text = redaction.redacted_text
                    
                    # Ingest into Atom Memory
                    if self.memory_handler:
                        from core.doc_freshness_service import (
                            extraction_content_hash,
                            extra_columns_for_ingest,
                        )

                        source_modified = file_info.get("modified_at")
                        content_hash = extraction_content_hash(text)
                        # Content-level idempotency: source modified_at can
                        # change (touch/rename) without the bytes changing —
                        # skip the rewrite and just refresh the cache marker.
                        if existing and existing.source_content_hash == content_hash:
                            existing.external_modified_at = source_modified
                            results["files_skipped"] += 1
                            continue
                        # Join-key bridge (hybrid search, Step 1): generate the PG
                        # row id BEFORE the LanceDB write and pass it as doc_id so
                        # the LanceDB documents row id equals the IngestedDocument
                        # id. Vector hits then resolve directly to
                        # documents.cat("knowledge/documents/<id>") paths.
                        new_id = f"doc_{datetime.now(timezone.utc).timestamp()}"
                        # Source URL: best-effort canonical locator from the
                        # integration-provided metadata.
                        source_url = (
                            file_info.get("url")
                            or file_info.get("web_url")
                            or file_info.get("webViewLink")
                            or f"{integration_id}:{file_info.get('path', '')}"
                        )

                        success = await asyncio.to_thread(
                            self.memory_handler.add_document,
                            table_name="documents",
                            text=text,
                            source=f"{integration_id}:{file_info.get('path', '')}",
                            metadata={
                                "file_name": file_info.get("name"),
                                "file_path": file_info.get("path"),
                                "file_type": file_ext,
                                "file_size": file_size,
                                "integration_id": integration_id,
                                "external_id": external_id,
                                "ingested_at": datetime.now(timezone.utc).isoformat(),
                                "source_url": source_url,
                                "source_content_hash": content_hash,
                                "freshness_status": "fresh",
                                "pg_document_id": new_id,
                                "source_type": "ingested",
                            },
                            user_id="system",
                            doc_id=new_id,
                            # Freshness columns as TOP-LEVEL filterable columns
                            # (not buried in the metadata JSON blob — see the
                            # warning in lancedb_handler.py add_document).
                            extra_columns=extra_columns_for_ingest(
                                freshness_status="fresh",
                                source_modified_at=source_modified
                                if isinstance(source_modified, datetime)
                                else None,
                                source_url=source_url,
                            ),
                        )

                        if success:
                            # Re-ingest of a modified file: remove the OLD
                            # vector row so search returns exactly one (fresh)
                            # copy instead of a stale+fresh duplicate pair.
                            if existing and existing.id and existing.id != new_id:
                                try:
                                    await asyncio.to_thread(
                                        self.memory_handler.delete_documents_by_id,
                                        "documents",
                                        existing.id,
                                    )
                                except Exception as old_del_err:  # noqa: BLE001 — best-effort cleanup
                                    logger.warning(
                                        f"Failed to remove superseded row {existing.id}: {old_del_err}"
                                    )
                            # Record ingestion (in-memory cache + DB for
                            # cross-run freshness tracking).
                            self.ingested_docs[external_id] = IngestedDocument(
                                id=new_id,
                                file_name=file_info.get("name", ""),
                                file_path=file_info.get("path", ""),
                                file_type=file_ext,
                                integration_id=integration_id,
                                workspace_id=self.workspace_id,
                                file_size_bytes=file_size,
                                content_preview=text[:500],
                                ingested_at=datetime.now(timezone.utc),
                                external_id=external_id,
                                external_modified_at=source_modified,
                                source_url=source_url,
                                source_content_hash=content_hash,
                                last_verified_at=datetime.now(timezone.utc),
                                source_modified_at=source_modified
                                if isinstance(source_modified, datetime)
                                else None,
                                freshness_status="fresh",
                            )
                            try:
                                self._persist_freshness_on_ingest(
                                    self.ingested_docs[external_id],
                                    source_url=source_url,
                                    content_hash=content_hash,
                                    source_modified_at=source_modified
                                    if isinstance(source_modified, datetime)
                                    else None,
                                )
                            except Exception as persist_err:
                                logger.warning(
                                    f"Freshness persist failed for {external_id}: {persist_err}"
                                )
                            results["files_ingested"] += 1
                            results["newly_ingested_files"].append(file_info.get("name"))

                            # Supersession: does this new doc obsolete an
                            # older same-topic doc in the same workspace?
                            try:
                                self._maybe_supersede_older_docs(
                                    text=text,
                                    new_doc_id=new_id,
                                    source_modified_at=source_modified,
                                )
                            except Exception as sup_err:
                                logger.warning(f"Supersession check failed for {new_id}: {sup_err}")
                
                except Exception as file_err:
                    results["errors"].append(f"{file_info.get('name')}: {str(file_err)}")

            # Reevaluate freshness across the workspace: tombstone docs that
            # were deleted upstream (absent from seen_external_ids) and age out
            # docs whose last verification is beyond the TTL.
            try:
                reeval = self._reevaluate_workspace(seen_external_ids)
                results["freshness"] = reeval
            except Exception as reeval_err:
                logger.warning(f"Freshness reevaluate failed for {integration_id}: {reeval_err}")
                results["freshness"] = {"error": str(reeval_err)}

            settings.last_sync = datetime.now(timezone.utc)
            self._persist_settings(settings)
            results["completed_at"] = datetime.now(timezone.utc).isoformat()
            results["success"] = True
            
            # TRIGGER AGENT IF FILES WERE INGESTED
            if results["files_ingested"] > 0:
                try:
                    from core.atom_meta_agent import handle_data_event_trigger
                    logger.info(f"Triggering Atom Agent for {results['files_ingested']} new documents in {integration_id}")

                    # Fire-and-forget (resolved: create_task, NOT await).
                    # Awaiting ran a FULL meta-agent turn inline — sync
                    # close-out blocked for minutes on LLM latency/retries.
                    # Mirrors the turn-fact extraction pending-set pattern:
                    # strong task ref prevents GC; done-callback discards.
                    _trigger_task = asyncio.create_task(handle_data_event_trigger(
                        event_type="document_ingestion",
                        data={
                            "integration_id": integration_id,
                            "count": results["files_ingested"],
                            "files": results["newly_ingested_files"]
                        },
                        workspace_id="default"
                    ))
                    _pending_agent_trigger_tasks.add(_trigger_task)
                    _trigger_task.add_done_callback(
                        lambda t: _pending_agent_trigger_tasks.discard(t)
                    )
                except Exception as trigger_err:
                    logger.warning(f"Failed to trigger agent after ingestion: {trigger_err}")
            
            
        except Exception as e:
            results["error"] = str(e)
            results["success"] = False
            logger.error(f"Sync failed for {integration_id}: {e}")

        return results

    # ====================================================================
    # Freshness helpers
    # ====================================================================

    def _freshness_session(self):
        """Open a short-lived DB session for freshness persistence.

        Mirrors the pattern in IngestionPipelineService._record_doc_ingestion
        (SessionLocal() opened and closed per call) so we don't leak a long-
        lived session across the async sync loop.
        """
        from core.database import SessionLocal
        return SessionLocal()

    def _mirror_pg_row(
        self,
        *,
        doc_id: str,
        workspace_id: str,
        file_name: str,
        file_path: str,
        file_type: str,
        integration_id: str,
        file_size_bytes: int,
        content_preview: str,
        external_id: str,
        content_hash: str,
        role: Optional[str] = None,
        source_modified_at: Optional[datetime] = None,
    ) -> None:
        """Upsert the IngestedDocument mirror row for a vector-first ingest.

        Keyed on ``id == doc_id`` — the same content-addressed/source-scoped
        identity used as the LanceDB doc_id — so hybrid search resolves the
        hit (bridged:true), the Knowledge VFS lists it under
        ``knowledge/documents/<id>``, and the FTS lexical leg indexes the
        preview. Idempotent across re-ingests: the identity contract in
        ``process_file_bytes`` yields a stable doc_id per (source, external_id)
        or per content hash.
        """
        from core.models import IngestedDocument as IngestedDocumentModel

        session = self._freshness_session()
        try:
            now = datetime.now(timezone.utc)
            row = (
                session.query(IngestedDocumentModel)
                .filter(IngestedDocumentModel.id == doc_id)
                .first()
            )
            if row is None:
                row = IngestedDocumentModel(
                    id=doc_id,
                    workspace_id=workspace_id or "default",
                    external_id=external_id,
                    integration_id=integration_id,
                )
                session.add(row)
            row.file_name = file_name
            row.file_path = file_path
            row.file_type = file_type
            row.integration_id = integration_id
            row.file_size_bytes = file_size_bytes
            row.content_preview = content_preview
            row.source_content_hash = content_hash
            row.last_verified_at = now
            row.ingested_at = now
            row.freshness_status = "fresh"
            # Update-detection inputs: freshness gates and hybrid-refresh
            # probes compare the SOURCE's modified time against these. Left
            # NULL, a vector-first ingest could never be detected as
            # out-of-date (the walker's modified_at check had nothing to
            # compare against).
            if source_modified_at is not None:
                row.source_modified_at = source_modified_at
                row.external_modified_at = source_modified_at
            if role:
                row.role = role
            session.commit()
        finally:
            session.close()

    def _persist_freshness_on_ingest(
        self,
        doc: IngestedDocument,
        *,
        source_url: Optional[str],
        content_hash: str,
        source_modified_at: Optional[datetime],
    ) -> None:
        """Upsert the IngestedDocument row and stamp it fresh.

        Persists to the ``ingested_documents`` table so freshness survives
        across sync runs (the in-memory ``self.ingested_docs`` cache is lost
        on restart). Also writes the new freshness columns.
        """
        from core.models import IngestedDocument as IngestedDocumentModel
        from core.doc_freshness_service import DocFreshnessService

        session = self._freshness_session()
        try:
            existing = (
                session.query(IngestedDocumentModel)
                .filter(
                    IngestedDocumentModel.workspace_id == doc.workspace_id,
                    IngestedDocumentModel.external_id == doc.external_id,
                )
                .first()
            )
            if existing is None:
                existing = IngestedDocumentModel(
                    id=doc.id,
                    workspace_id=doc.workspace_id,
                    file_name=doc.file_name,
                    file_path=doc.file_path,
                    file_type=doc.file_type,
                    integration_id=doc.integration_id,
                    file_size_bytes=doc.file_size_bytes,
                    content_preview=doc.content_preview,
                    external_id=doc.external_id,
                    external_modified_at=doc.external_modified_at,
                )
                session.add(existing)
            else:
                if existing.id != doc.id:
                    # Join-key realignment: the caller rewrote the vector row
                    # under doc.id and deleted the row under existing.id, so
                    # keeping the old PG id would leave cat/search resolving to
                    # a deleted vector row and the new one bridged:false.
                    session.delete(existing)
                    session.flush()
                    existing = IngestedDocumentModel(
                        id=doc.id,
                        workspace_id=doc.workspace_id,
                        file_name=doc.file_name,
                        file_path=doc.file_path,
                        file_type=doc.file_type,
                        integration_id=doc.integration_id,
                        file_size_bytes=doc.file_size_bytes,
                        content_preview=doc.content_preview,
                        external_id=doc.external_id,
                        external_modified_at=doc.external_modified_at,
                    )
                    session.add(existing)
                else:
                    # Keep columns in sync with the freshly fetched version.
                    existing.file_name = doc.file_name
                    existing.file_path = doc.file_path
                    existing.file_type = doc.file_type
                    existing.file_size_bytes = doc.file_size_bytes
                    existing.content_preview = doc.content_preview
                    existing.external_modified_at = doc.external_modified_at

            svc = DocFreshnessService(session, workspace_id=doc.workspace_id)
            svc.mark_on_ingest(
                existing,
                source_url=source_url,
                content_hash=content_hash,
                source_modified_at=source_modified_at,
            )
        finally:
            session.close()

    def _mark_doc_stale(self, doc: IngestedDocument, *, reason: str) -> None:
        """Record a doc as stale (source changed) before re-ingest."""
        from core.models import IngestedDocument as IngestedDocumentModel
        from core.doc_freshness_service import DocFreshnessService

        doc.freshness_status = "stale"
        session = self._freshness_session()
        try:
            row = (
                session.query(IngestedDocumentModel)
                .filter(
                    IngestedDocumentModel.workspace_id == doc.workspace_id,
                    IngestedDocumentModel.external_id == doc.external_id,
                )
                .first()
            )
            if row is not None:
                svc = DocFreshnessService(session, workspace_id=doc.workspace_id)
                svc.mark_stale(row, reason=reason)
        finally:
            session.close()

    def _reevaluate_workspace(self, seen_external_ids: Set[str]) -> Dict[str, Any]:
        """Recompute freshness for all docs in this workspace.

        Returns the summary dict from DocFreshnessService.reevaluate_workspace.
        """
        from core.doc_freshness_service import DocFreshnessService

        session = self._freshness_session()
        try:
            svc = DocFreshnessService(session, workspace_id=self.workspace_id)
            summary = svc.reevaluate_workspace(self.workspace_id, seen_external_ids)
            return summary.as_dict()
        finally:
            session.close()

    def _maybe_supersede_older_docs(
        self,
        *,
        text: str,
        new_doc_id: str,
        source_modified_at: Optional[datetime],
    ) -> None:
        """Detect whether this newly ingested doc supersedes older same-topic docs.

        Hybrid detection (see doc_freshness_service.detect_supersession):
        semantic near-duplicate OR entity overlap, confirmed by a newer-
        timestamp heuristic. Candidates are marked ``superseded`` and linked,
        and the Postgres GraphRAG cascade stamps their derived nodes/edges.
        """
        from core.doc_freshness_service import (
            DocFreshnessService,
            detect_supersession,
            doc_ts,
        )
        from core.models import IngestedDocument as IngestedDocumentModel

        session = self._freshness_session()
        try:
            # Gather older docs in the same workspace that are still
            # fresh/stale (candidates for supersession). Bounded by recent docs
            # to keep the per-ingest cost predictable.
            older_rows = (
                session.query(IngestedDocumentModel)
                .filter(
                    IngestedDocumentModel.workspace_id == self.workspace_id,
                    IngestedDocumentModel.id != new_doc_id,
                    IngestedDocumentModel.freshness_status.in_(["fresh", "stale"]),
                )
                .order_by(IngestedDocumentModel.ingested_at.desc())
                .limit(50)
                .all()
            )
            if not older_rows:
                return

            older_docs = [
                {
                    "doc_id": r.id,
                    "text": r.content_preview or "",
                    "ingested_at": r.ingested_at,
                    "external_modified_at": r.external_modified_at or r.source_modified_at,
                    "freshness_status": r.freshness_status,
                }
                for r in older_rows
            ]

            # Newer doc's embedding + entity set.
            newer_embedding = None
            if self.memory_handler is not None:
                try:
                    newer_embedding = self.memory_handler.embed_text(text)
                    if newer_embedding is not None and hasattr(newer_embedding, "tolist"):
                        newer_embedding = newer_embedding.tolist()
                except Exception:
                    newer_embedding = None

            svc = DocFreshnessService(session, workspace_id=self.workspace_id)
            newer_entities = svc.entity_set_for_doc(new_doc_id)
            older_entity_sets = {r.id: svc.entity_set_for_doc(r.id) for r in older_rows}

            candidates = detect_supersession(
                newer_doc_id=new_doc_id,
                newer_text=text,
                newer_embedding=newer_embedding,
                newer_entities=newer_entities,
                newer_ts=source_modified_at if isinstance(source_modified_at, datetime) else None,
                older_docs=older_docs,
                older_entity_sets=older_entity_sets,
            )
            if not candidates:
                return

            logger.info(
                f"Supersession: {new_doc_id} obsoletes {len(candidates)} older doc(s)"
            )
            svc.apply_supersession(
                candidates,
                new_doc_id,
                cascade_to_graph=svc.cascade_graph_supersession,
            )
        finally:
            session.close()

    async def _list_files(
        self, 
        integration_id: str, 
        settings: IngestionSettings
    ) -> List[Dict[str, Any]]:
        """List files from an integration"""
        files = []
        
        try:
            if integration_id == "google_drive":
                files = await self._list_google_drive_files(settings)
            elif integration_id == "dropbox":
                files = await self._list_dropbox_files(settings)
            elif integration_id == "onedrive":
                files = await self._list_onedrive_files(settings)
            elif integration_id == "notion":
                files = await self._list_notion_pages(settings)
            else:
                logger.warning(f"No file lister for {integration_id}")
        
        except Exception as e:
            logger.error(f"Failed to list files from {integration_id}: {e}")
        
        return files
    
    async def _download_file(
        self, 
        integration_id: str, 
        file_info: Dict[str, Any]
    ) -> Optional[bytes]:
        """Download file content from an integration"""
        try:
            if integration_id == "google_drive":
                return await self._download_google_drive_file(file_info)
            elif integration_id == "dropbox":
                return await self._download_dropbox_file(file_info)
            elif integration_id == "onedrive":
                return await self._download_onedrive_file(file_info)
            elif integration_id == "notion":
                return await self._download_notion_content(file_info)
            else:
                logger.warning(f"No downloader for {integration_id}")
                return None
        
        except Exception as e:
            logger.error(f"Failed to download from {integration_id}: {e}")
            return None
    
    # Integration-specific implementations
    async def _list_google_drive_files(self, settings: IngestionSettings) -> List[Dict]:
        """List files from Google Drive"""
        try:
            import os

            from integrations.google_drive_service import google_drive_service

            access_token = os.getenv("GOOGLE_DRIVE_ACCESS_TOKEN")
            if not access_token:
                logger.warning("Google Drive access token not configured")
                return []

            result = await google_drive_service.list_files(
                access_token=access_token,
                page_size=getattr(settings, "max_files", 100)
            )

            if result["status"] == "success":
                files = result["data"].get("files", [])
                logger.info(f"Listed {len(files)} files from Google Drive")
                return files
            else:
                logger.error(f"Failed to list Google Drive files: {result.get('message')}")
                return []

        except Exception as e:
            logger.error(f"Google Drive file listing error: {e}")
            return []

    async def _download_google_drive_file(self, file_info: Dict) -> Optional[bytes]:
        """Download from Google Drive"""
        try:
            import os
            import httpx

            access_token = os.getenv("GOOGLE_DRIVE_ACCESS_TOKEN")
            if not access_token:
                logger.warning("Google Drive access token not configured")
                return None

            file_id = file_info.get("id")
            if not file_id:
                return None

            # Get download URL based on MIME type
            mime_type = file_info.get("mimeType", "")
            if "google-apps" in mime_type:
                # Google Docs format - need to export
                export_formats = {
                    "application/vnd.google-apps.document": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "application/vnd.google-apps.spreadsheet": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    "application/vnd.google-apps.presentation": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                }
                export_mime = export_formats.get(mime_type, "application/pdf")
                url = f"https://www.googleapis.com/drive/v3/files/{file_id}/export?mimeType={export_mime}"
            else:
                # Regular file - direct download
                url = f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media"

            async with httpx.AsyncClient() as client:
                response = await client.get(
                    url,
                    headers={"Authorization": f"Bearer {access_token}"},
                    timeout=60.0
                )
                response.raise_for_status()
                return response.content

        except Exception as e:
            logger.error(f"Google Drive download error: {e}")
            return None
    
    async def _list_dropbox_files(self, settings: IngestionSettings) -> List[Dict]:
        """List files from Dropbox"""
        try:
            import os
            import httpx

            access_token = os.getenv("DROPBOX_ACCESS_TOKEN")
            if not access_token:
                logger.warning("Dropbox access token not configured")
                return []

            # List files in Dropbox using API v2
            async with httpx.AsyncClient() as client:
                headers = {
                    "Authorization": f"Bearer {access_token}",
                    "Content-Type": "application/json"
                }

                # Use list_folder endpoint
                response = await client.post(
                    "https://api.dropboxapi.com/2/files/list_folder",
                    headers=headers,
                    json={"path": "", "recursive": False, "limit": getattr(settings, "max_files", 100)},
                    timeout=30.0
                )
                response.raise_for_status()
                data = response.json()

                files = []
                for entry in data.get("entries", []):
                    if entry.get(".tag") == "file":
                        files.append({
                            "id": entry.get("id"),
                            "name": entry.get("name"),
                            "path_lower": entry.get("path_lower"),
                            "size": entry.get("size"),
                            "client_modified": entry.get("client_modified"),
                            "server_modified": entry.get("server_modified")
                        })

                logger.info(f"Listed {len(files)} files from Dropbox")
                return files

        except Exception as e:
            logger.error(f"Dropbox file listing error: {e}")
            return []

    async def _download_dropbox_file(self, file_info: Dict) -> Optional[bytes]:
        """Download from Dropbox"""
        try:
            import os
            import httpx

            access_token = os.getenv("DROPBOX_ACCESS_TOKEN")
            if not access_token:
                logger.warning("Dropbox access token not configured")
                return None

            path = file_info.get("path_lower")
            if not path:
                return None

            # Download file using Dropbox API
            async with httpx.AsyncClient() as client:
                headers = {
                    "Authorization": f"Bearer {access_token}",
                    "Content-Type": "application/json"
                }

                # Get temporary download link
                post_response = await client.post(
                    "https://api.dropboxapi.com/2/files/get_temporary_link",
                    headers=headers,
                    json={"path": path},
                    timeout=30.0
                )
                post_response.raise_for_status()
                link_data = post_response.json()

                # Download file content
                download_url = link_data.get("link")
                if download_url:
                    download_response = await client.get(download_url, timeout=60.0)
                    download_response.raise_for_status()
                    return download_response.content

        except Exception as e:
            logger.error(f"Dropbox download error: {e}")
            return None
    
    async def _list_onedrive_files(self, settings: IngestionSettings) -> List[Dict]:
        """List files from OneDrive via the Graph-backed integration service."""
        try:
            import os

            from integrations.onedrive_service import onedrive_service

            access_token = os.getenv("ONEDRIVE_ACCESS_TOKEN")
            if not access_token:
                logger.warning("OneDrive access token not configured")
                return []

            result = await onedrive_service.list_files(
                access_token=access_token,
                page_size=getattr(settings, "max_files", 100),
            )
            if result.get("status") != "success":
                logger.error(f"Failed to list OneDrive files: {result.get('message')}")
                return []

            files = []
            for item in result.get("data", {}).get("value", []):
                # Skip folders — only file items have a "file" facet.
                if "file" not in (item or {}):
                    continue
                files.append({
                    "id": item.get("id"),
                    "name": item.get("name"),
                    "path": item.get("parentReference", {}).get("path", ""),
                    "size": item.get("size", 0),
                    "modified_at": item.get("lastModifiedDateTime"),
                    "url": item.get("webUrl"),
                })
            logger.info(f"Listed {len(files)} files from OneDrive")
            return files

        except Exception as e:
            logger.error(f"OneDrive file listing error: {e}")
            return []

    async def _download_onedrive_file(self, file_info: Dict) -> Optional[bytes]:
        """Download from OneDrive via the Graph-backed integration service."""
        try:
            import os

            from integrations.onedrive_service import onedrive_service

            access_token = os.getenv("ONEDRIVE_ACCESS_TOKEN")
            if not access_token:
                logger.warning("OneDrive access token not configured")
                return None

            file_id = file_info.get("id")
            if not file_id:
                return None
            return await onedrive_service.download_file_bytes(
                access_token=access_token, file_id=file_id
            )

        except Exception as e:
            logger.error(f"OneDrive download error: {e}")
            return None

    async def _list_notion_pages(self, settings: IngestionSettings) -> List[Dict]:
        """List pages from Notion via the search endpoint."""
        try:
            import os

            from integrations.notion_service import NotionService

            token = os.getenv("NOTION_ACCESS_TOKEN") or os.getenv("NOTION_API_KEY") or os.getenv("NOTION_TOKEN")
            if not token:
                logger.warning("Notion access token not configured")
                return []

            notion = NotionService(config={"access_token": token})
            data = notion.search(page_size=getattr(settings, "max_files", 100))
            pages = []
            for item in data.get("results", []):
                if item.get("object") != "page":
                    continue
                # Best-effort title: check common title property locations.
                title = item.get("url", "").rsplit("/", 1)[-1] or "Untitled"
                props = item.get("properties", {})
                for value in props.values():
                    if value.get("type") == "title" and value.get("title"):
                        title = "".join(
                            part.get("plain_text", "") for part in value["title"]
                        )
                        break
                pages.append({
                    "id": item.get("id"),
                    "name": f"{title}.md",
                    "path": title,
                    "size": 0,
                    "modified_at": item.get("last_edited_time"),
                    "url": item.get("url"),
                })
            logger.info(f"Listed {len(pages)} pages from Notion")
            return pages

        except Exception as e:
            logger.error(f"Notion page listing error: {e}")
            return []

    async def _download_notion_content(self, file_info: Dict) -> Optional[bytes]:
        """Flatten a Notion page's blocks into markdown-ish text bytes."""
        try:
            import os

            from integrations.notion_service import NotionService

            token = os.getenv("NOTION_ACCESS_TOKEN") or os.getenv("NOTION_API_KEY") or os.getenv("NOTION_TOKEN")
            if not token:
                logger.warning("Notion access token not configured")
                return None

            page_id = file_info.get("id")
            if not page_id:
                return None

            notion = NotionService(config={"access_token": token})
            parts: List[str] = [f"# {file_info.get('path') or file_info.get('name') or page_id}"]
            cursor = None
            while True:
                kwargs = {"page_size": 100}
                if cursor:
                    kwargs["start_cursor"] = cursor
                data = notion.get_block_children(page_id, **kwargs)
                results = data.get("results", [])
                if not results and not cursor:
                    break
                for block in results:
                    # Any rich-text-bearing sub-key of the block payload.
                    for value in block.values():
                        if isinstance(value, dict) and isinstance(value.get("rich_text"), list):
                            text = "".join(
                                rt.get("plain_text", "")
                                for rt in value["rich_text"]
                            )
                            if text:
                                prefix = "- " if block.get("type") == "bulleted_list_item" else ""
                                parts.append(prefix + text)
                        elif isinstance(value, dict) and isinstance(value.get("title"), list):
                            text = "".join(
                                rt.get("plain_text", "") for rt in value["title"]
                            )
                            if text:
                                parts.append(text)
                cursor = data.get("next_cursor")
                if not data.get("has_more") or not cursor:
                    break

            return "\n\n".join(parts).encode("utf-8")

        except Exception as e:
            logger.error(f"Notion content download error: {e}")
            return None
    
    def get_ingested_documents(
        self, 
        integration_id: Optional[str] = None,
        file_type: Optional[str] = None
    ) -> List[IngestedDocument]:
        """Get list of ingested documents"""
        docs = list(self.ingested_docs.values())
        
        if integration_id:
            docs = [d for d in docs if d.integration_id == integration_id]
        if file_type:
            docs = [d for d in docs if d.file_type == file_type]
        
        return docs
    
    async def remove_integration_documents(
        self, 
        integration_id: str
    ) -> Dict[str, Any]:
        """
        Remove all ingested documents from a specific integration.
        Clears from Atom Memory (LanceDB + GraphRAG).
        """
        count = 0
        removed_ids = []

        for ext_id, doc in list(self.ingested_docs.items()):
            if doc.integration_id == integration_id:
                removed_ids.append(ext_id)
                # Real vector cleanup: delete the stored row by its doc id.
                if self.memory_handler and doc.id:
                    try:
                        await asyncio.to_thread(
                            self.memory_handler.delete_documents_by_id,
                            "documents",
                            doc.id,
                        )
                    except Exception as del_err:  # noqa: BLE001 — removal best-effort
                        logger.warning(f"LanceDB delete failed for {doc.id}: {del_err}")
                del self.ingested_docs[ext_id]
                count += 1

        # Best-effort removal of the durable freshness rows so deleted docs
        # don't reappear in freshness reevaluations.
        try:
            from core.models import IngestedDocument as IngestedDocumentRow

            with self._freshness_session() as db:
                db.query(IngestedDocumentRow).filter(
                    IngestedDocumentRow.integration_id == integration_id
                ).delete(synchronize_session=False)
                db.commit()
        except Exception as db_err:  # noqa: BLE001
            logger.warning(f"IngestedDocument cleanup skipped for {integration_id}: {db_err}")

        logger.info(f"Removed {count} documents from {integration_id}")
        
        return {
            "success": True,
            "integration_id": integration_id,
            "documents_removed": count,
            "removed_ids": removed_ids
        }
    
    def get_all_settings(self) -> List[Dict[str, Any]]:
        """Get settings for all integrations"""
        return [
            {
                "integration_id": s.integration_id,
                "enabled": s.enabled,
                "auto_sync_new_files": s.auto_sync_new_files,
                "file_types": s.file_types,
                "sync_folders": s.sync_folders,
                "max_file_size_mb": s.max_file_size_mb,
                "sync_frequency_minutes": s.sync_frequency_minutes,
                "last_sync": s.last_sync.isoformat() if s.last_sync else None,
            }
            for s in self.settings.values()
        ]


# Per-workspace singleton registry (R84d). ``_doc_ingestion_service`` is kept
# as the "default"-workspace alias for backward compatibility.
_doc_ingestion_service: Optional[AutoDocumentIngestionService] = None
_doc_ingestion_services: Dict[str, AutoDocumentIngestionService] = {}


def get_document_ingestion_service(
    workspace_id: str = "default",
) -> AutoDocumentIngestionService:
    """Get or create the document ingestion service for a workspace.

    One instance per workspace: each binds its own LanceDB handler, settings
    cache and durable ``ingestion_settings`` rows to that workspace. Callers
    that pass no workspace get the shared "default" instance (previous
    behavior).
    """
    global _doc_ingestion_service
    ws = (workspace_id or "default").strip() or "default"
    service = _doc_ingestion_services.get(ws)
    if service is None:
        service = AutoDocumentIngestionService(workspace_id=ws)
        _doc_ingestion_services[ws] = service
        if ws == "default":
            _doc_ingestion_service = service
    return service


def reset_document_ingestion_services() -> None:
    """Drop all cached per-workspace instances (test isolation helper)."""
    global _doc_ingestion_service
    _doc_ingestion_services.clear()
    _doc_ingestion_service = None


# Alias for backward compatibility with tests
AutoDocumentIngestion = AutoDocumentIngestionService
