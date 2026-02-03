import json
import logging
import os
from typing import Any, Dict, List
import pandas as pd
from docx import Document

from core.business_intelligence import BusinessEventIntelligence
from core.knowledge_extractor import KnowledgeExtractor

logger = logging.getLogger(__name__)

import PyPDF2


class DocumentLifecycleLearner:
    """
    Extracts business lifecycle events from Excel, Word, and PDF documents.
    """

    def __init__(self, ai_service: Any = None, db_session: Any = None):
        self.extractor = KnowledgeExtractor(ai_service)
        self.biz_intel = BusinessEventIntelligence(db_session)

    async def learn_from_file(self, file_path: str, workspace_id: str):
        """
        Main entry point for learning from a local file.
        """
        ext = os.path.splitext(file_path)[1].lower()
        content = ""
        
        if ext in [".xlsx", ".xls", ".csv"]:
            content = self._parse_excel(file_path)
        elif ext == ".docx":
            content = self._parse_word(file_path)
        elif ext == ".pdf":
            content = self._parse_pdf(file_path)
        else:
            logger.warning(f"Unsupported file format: {ext}")
            return
            
        if content:
            # Pass to extractor
            knowledge = await self.extractor.extract_knowledge(content, source=f"document_{os.path.basename(file_path)}")
            await self.biz_intel.process_extracted_events(knowledge, workspace_id)

    def _parse_excel(self, file_path: str) -> str:
        """Parses all sheets of an Excel file into a text representation."""
        try:
            if file_path.endswith(".csv"):
                df = pd.read_csv(file_path)
                return df.to_string()
            else:
                # Read all sheets
                xls = pd.ExcelFile(file_path)
                full_text = []
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    full_text.append(f"--- Sheet: {sheet_name} ---")
                    full_text.append(df.to_string())
                return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Failed to parse Excel {file_path}: {e}")
            return ""

    def _parse_word(self, file_path: str) -> str:
        """Parses Word document into text."""
        try:
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))
                    
            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Failed to parse Word doc {file_path}: {e}")
            return ""

    def _parse_pdf(self, file_path: str) -> str:
        """Parses PDF document into text."""
        try:
            text = []
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text())
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}")
            return ""
