"""
Content Extractor - Multi-format Content Processing Pipeline
Extracts text, metadata, and insights from various file formats
"""

import os
import json
import asyncio
import logging
import hashlib
import mimetypes
from datetime import datetime
from typing import Dict, List, Optional, Any, Union, Tuple
from pathlib import Path
import aiofiles
from dataclasses import dataclass, asdict
from contextlib import asynccontextmanager

# Content processing imports
try:
    import PyPDF2
    import fitz  # PyMuPDF
    import pdfplumber
    import docx
    import python_pptx
    from PIL import Image
    import cv2
    import pytesseract
    from bs4 import BeautifulSoup
    import markdown
    import textract
    TEXT_PROCESSING_AVAILABLE = True
except ImportError:
    TEXT_PROCESSING_AVAILABLE = False
    logging.warning("Text processing libraries not available")

# Local imports
from loguru import logger
from config import get_config_instance

@dataclass
class ExtractedContent:
    """Extracted content data model"""
    file_id: str
    file_name: str
    file_path: str
    mime_type: str
    file_size: int
    
    # Extracted content
    text_content: str
    html_content: str
    markdown_content: str
    
    # Metadata
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    keywords: List[str] = None
    language: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    character_count: Optional[int] = None
    
    # File-specific metadata
    image_metadata: Dict[str, Any] = None
    document_metadata: Dict[str, Any] = None
    video_metadata: Dict[str, Any] = None
    audio_metadata: Dict[str, Any] = None
    
    # Processing info
    extraction_method: str = ""
    processing_time: float = 0.0
    success: bool = True
    error_message: Optional[str] = None
    
    # Timestamps
    created_at: datetime = None
    processed_at: datetime = None
    
    def __post_init__(self):
        if self.keywords is None:
            self.keywords = []
        if self.image_metadata is None:
            self.image_metadata = {}
        if self.document_metadata is None:
            self.document_metadata = {}
        if self.video_metadata is None:
            self.video_metadata = {}
        if self.audio_metadata is None:
            self.audio_metadata = {}
        if self.created_at is None:
            self.created_at = datetime.utcnow()
        if self.processed_at is None:
            self.processed_at = datetime.utcnow()
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return asdict(self)
    
    @property
    def content_hash(self) -> str:
        """Get content hash for deduplication"""
        content = f"{self.text_content}{self.title}{self.author}"
        return hashlib.sha256(content.encode()).hexdigest()
    
    @property
    def summary(self) -> str:
        """Get content summary"""
        if self.title and self.text_content:
            return f"{self.title}: {self.text_content[:200]}..."
        elif self.title:
            return self.title
        elif self.text_content:
            return self.text_content[:200] + "..."
        return "No content available"

class ContentExtractor:
    """Multi-format Content Extractor"""
    
    def __init__(self, config=None):
        self.config = config or get_config_instance()
        self.ingestion_config = self.config.ingestion
        
        if not TEXT_PROCESSING_AVAILABLE:
            logger.warning("Text processing libraries not available. Some features may not work.")
        
        # Processing configuration
        self.max_file_size = self.ingestion_config.max_file_size
        self.supported_formats = self.ingestion_config.supported_formats
        self.enable_ocr = self.ingestion_config.enable_ocr
        self.enable_metadata_extraction = self.ingestion_config.enable_metadata_extraction
        self.extraction_timeout = self.ingestion_config.extraction_timeout
        
        # File format handlers
        self.format_handlers = {
            # Documents
            '.pdf': self._extract_pdf,
            '.doc': self._extract_doc,
            '.docx': self._extract_docx,
            '.txt': self._extract_text,
            '.rtf': self._extract_rtf,
            '.odt': self._extract_odt,
            
            # Presentations
            '.ppt': self._extract_ppt,
            '.pptx': self._extract_pptx,
            
            # Spreadsheets
            '.xls': self._extract_xls,
            '.xlsx': self._extract_xlsx',
            '.csv': self._extract_csv',
            
            # Web content
            '.html': self._extract_html,
            '.htm': self._extract_html,
            '.xml': self._extract_xml,
            '.md': self._extract_markdown,
            
            # Images
            '.jpg': self._extract_image,
            '.jpeg': self._extract_image,
            '.png': self._extract_image,
            '.gif': self._extract_image,
            '.bmp': self._extract_image,
            '.tiff': self._extract_image,
            '.svg': self._extract_svg,
            
            # Archives
            '.zip': self._extract_archive,
            '.rar': self._extract_archive,
            '.7z': self._extract_archive,
            '.tar': self._extract_archive,
            '.gz': self._extract_archive
        }
        
        logger.info("Content Extractor initialized")
    
    async def extract_content(self, 
                            file_id: str,
                            file_path: str,
                            file_name: Optional[str] = None) -> ExtractedContent:
        """Extract content from file"""
        
        start_time = asyncio.get_event_loop().time()
        
        try:
            # Validate file
            validation_result = await self._validate_file(file_path, file_name)
            if not validation_result["valid"]:
                return ExtractedContent(
                    file_id=file_id,
                    file_name=file_name or Path(file_path).name,
                    file_path=file_path,
                    mime_type="",
                    file_size=0,
                    text_content="",
                    html_content="",
                    markdown_content="",
                    success=False,
                    error_message=validation_result["error"],
                    processing_time=asyncio.get_event_loop().time() - start_time
                )
            
            file_info = validation_result["file_info"]
            
            # Get file handler
            file_ext = Path(file_info["file_name"]).suffix.lower()
            handler = self.format_handlers.get(file_ext)
            
            if not handler:
                # Try generic text extraction
                handler = self._extract_generic_text
            
            # Extract content
            if asyncio.iscoroutinefunction(handler):
                content_result = await handler(file_info["file_path"])
            else:
                content_result = await asyncio.get_event_loop().run_in_executor(
                    None,
                    handler,
                    file_info["file_path"]
                )
            
            # Create extracted content object
            extracted_content = ExtractedContent(
                file_id=file_id,
                file_name=file_info["file_name"],
                file_path=file_info["file_path"],
                mime_type=file_info["mime_type"],
                file_size=file_info["file_size"],
                text_content=content_result.get("text_content", ""),
                html_content=content_result.get("html_content", ""),
                markdown_content=content_result.get("markdown_content", ""),
                title=content_result.get("title"),
                author=content_result.get("author"),
                subject=content_result.get("subject"),
                keywords=content_result.get("keywords", []),
                language=content_result.get("language"),
                page_count=content_result.get("page_count"),
                word_count=content_result.get("word_count"),
                character_count=content_result.get("character_count"),
                image_metadata=content_result.get("image_metadata", {}),
                document_metadata=content_result.get("document_metadata", {}),
                video_metadata=content_result.get("video_metadata", {}),
                audio_metadata=content_result.get("audio_metadata", {}),
                extraction_method=content_result.get("extraction_method", "unknown"),
                success=content_result.get("success", True),
                error_message=content_result.get("error_message"),
                processing_time=asyncio.get_event_loop().time() - start_time
            )
            
            # Count words and characters
            if extracted_content.text_content:
                words = extracted_content.text_content.split()
                extracted_content.word_count = len(words)
                extracted_content.character_count = len(extracted_content.text_content)
            
            logger.info(f"Extracted content from {file_name}: {extracted_content.processing_time:.2f}s")
            
            return extracted_content
        
        except Exception as e:
            logger.error(f"Failed to extract content from {file_name}: {e}")
            return ExtractedContent(
                file_id=file_id,
                file_name=file_name or Path(file_path).name,
                file_path=file_path,
                mime_type="",
                file_size=0,
                text_content="",
                html_content="",
                markdown_content="",
                success=False,
                error_message=str(e),
                processing_time=asyncio.get_event_loop().time() - start_time
            )
    
    # ==================== FILE VALIDATION ====================
    
    async def _validate_file(self, file_path: str, file_name: Optional[str] = None) -> Dict[str, Any]:
        """Validate file for processing"""
        
        try:
            file_path_obj = Path(file_path)
            
            # Check if file exists
            if not file_path_obj.exists():
                return {
                    "valid": False,
                    "error": "File does not exist"
                }
            
            # Get file info
            file_stat = file_path_obj.stat()
            file_size = file_stat.st_size
            actual_file_name = file_name or file_path_obj.name
            
            # Check file size
            if file_size > self.max_file_size:
                return {
                    "valid": False,
                    "error": f"File too large: {file_size} bytes (max: {self.max_file_size} bytes)"
                }
            
            # Get MIME type
            mime_type, _ = mimetypes.guess_type(file_path)
            
            # Check if format is supported
            file_ext = file_path_obj.suffix.lower()
            if file_ext not in self.supported_formats:
                return {
                    "valid": False,
                    "error": f"Unsupported file format: {file_ext}"
                }
            
            return {
                "valid": True,
                "file_info": {
                    "file_name": actual_file_name,
                    "file_path": str(file_path_obj),
                    "mime_type": mime_type or "application/octet-stream",
                    "file_size": file_size,
                    "file_extension": file_ext
                }
            }
        
        except Exception as e:
            logger.error(f"File validation failed: {e}")
            return {
                "valid": False,
                "error": f"Validation error: {str(e)}"
            }
    
    # ==================== DOCUMENT EXTRACTORS ====================
    
    def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract content from PDF file"""
        
        try:
            content = {
                "text_content": "",
                "html_content": "",
                "markdown_content": "",
                "extraction_method": "pdf",
                "success": True
            }
            
            # Try PyMuPDF first (faster and more accurate)
            try:
                import fitz
                doc = fitz.open(file_path)
                text_parts = []
                
                metadata = doc.metadata
                content.update({
                    "title": metadata.get("title"),
                    "author": metadata.get("author"),
                    "subject": metadata.get("subject"),
                    "page_count": doc.page_count,
                    "document_metadata": metadata
                })
                
                for page in doc:
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                
                content["text_content"] = "\n\n".join(text_parts)
                
                # Create HTML
                html_parts = []
                for i, page in enumerate(doc):
                    page_html = f"<div class='pdf-page' data-page='{i+1}'>\n"
                    page_html += f"<pre>{page.get_text()}</pre>\n"
                    page_html += "</div>"
                    html_parts.append(page_html)
                
                content["html_content"] = f"<div class='pdf-document'>\n" + "\n".join(html_parts) + "\n</div>"
                
                doc.close()
            
            except ImportError:
                # Fall back to PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    content.update({
                        "page_count": len(pdf_reader.pages),
                        "document_metadata": pdf_reader.metadata or {}
                    })
                    
                    if pdf_reader.metadata:
                        content.update({
                            "title": pdf_reader.metadata.get("/Title"),
                            "author": pdf_reader.metadata.get("/Author"),
                            "subject": pdf_reader.metadata.get("/Subject")
                        })
                    
                    text_parts = []
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_parts.append(page_text)
                    
                    content["text_content"] = "\n\n".join(text_parts)
                    
                    # Create HTML
                    html_parts = []
                    for i, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        page_html = f"<div class='pdf-page' data-page='{i+1}'>\n"
                        page_html += f"<pre>{page_text}</pre>\n"
                        page_html += "</div>"
                        html_parts.append(page_html)
                    
                    content["html_content"] = f"<div class='pdf-document'>\n" + "\n".join(html_parts) + "\n</div>"
            
            # Create markdown
            content["markdown_content"] = self._html_to_markdown(content["html_content"])
            
            return content
        
        except Exception as e:
            logger.error(f"Failed to extract PDF content: {e}")
            return {
                "text_content": "",
                "html_content": "",
                "markdown_content": "",
                "extraction_method": "pdf",
                "success": False,
                "error_message": str(e)
            }
    
    def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract content from DOCX file"""
        
        try:
            doc = docx.Document(file_path)
            
            content = {
                "text_content": "",
                "html_content": "",
                "markdown_content": "",
                "extraction_method": "docx",
                "success": True
            }
            
            # Extract text
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            content["text_content"] = "\n\n".join(text_parts)
            
            # Extract metadata
            core_props = doc.core_properties
            content.update({
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "keywords": core_props.keywords.split(";") if core_props.keywords else [],
                "document_metadata": {
                    "title": core_props.title,
                    "author": core_props.author,
                    "subject": core_props.subject,
                    "keywords": core_props.keywords,
                    "category": core_props.category,
                    "comments": core_props.comments,
                    "created": core_props.created.isoformat() if core_props.created else None,
                    "modified": core_props.modified.isoformat() if core_props.modified else None,
                    "last_modified_by": core_props.last_modified_by,
                    "revision": core_props.revision
                }
            })
            
            # Create HTML
            html_parts = []
            for paragraph in doc.paragraphs:
                style = paragraph.style.name
                text = paragraph.text.strip()
                
                if text:
                    if style.startswith("Heading"):
                        level = style.split()[-1]
                        html_parts.append(f"<h{level}>{text}</h{level}>")
                    else:
                        html_parts.append(f"<p>{text}</p>")
            
            content["html_content"] = f"<div class='docx-document'>\n" + "\n".join(html_parts) + "\n</div>"
            
            # Create markdown
            content["markdown_content"] = self._html_to_markdown(content["html_content"])
            
            return content
        
        except Exception as e:
            logger.error(f"Failed to extract DOCX content: {e}")
            return {
                "text_content": "",
                "html_content": "",
                "markdown_content": "",
                "extraction_method": "docx",
                "success": False,
                "error_message": str(e)
            }
    
    def _extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract content from plain text file"""
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text_content = file.read()
            
            # Create HTML
            html_content = f"<div class='text-document'>\n<pre>{text_content}</pre>\n</div>"
            
            # Create markdown (text is already markdown-compatible)
            markdown_content = text_content
            
            return {
                "text_content": text_content,
                "html_content": html_content,
                "markdown_content": markdown_content,
                "extraction_method": "text",
                "success": True
            }
        
        except Exception as e:
            logger.error(f"Failed to extract text content: {e}")
            return {
                "text_content": "",
                "html_content": "",
                "markdown_content": "",
                "extraction_method": "text",
                "success": False,
                "error_message": str(e)
            }
    
    # ==================== IMAGE EXTRACTORS ====================
    
    def _extract_image(self, file_path: str) -> Dict[str, Any]:
        """Extract content from image file"""
        
        try:
            # Extract metadata
            image = Image.open(file_path)
            
            content = {
                "text_content": "",
                "html_content": "",
                "markdown_content": "",
                "extraction_method": "image",
                "success": True,
                "image_metadata": {
                    "format": image.format,
                    "mode": image.mode,
                    "size": image.size,
                    "width": image.size[0],
                    "height": image.size[1],
                    "has_transparency": image.mode in ('RGBA', 'LA'),
                    "color_palette": image.getpalette() if image.mode == 'P' else None
                }
            }
            
            # Extract EXIF data if available
            try:
                from PIL.ExifTags import TAGS
                exif_data = {}
                
                if hasattr(image, '_getexif') and image._getexif() is not None:
                    exif = image._getexif()
                    for tag_id, value in exif.items():
                        tag = TAGS.get(tag_id, tag_id)
                        exif_data[tag] = value
                
                content["image_metadata"]["exif"] = exif_data
            
            except Exception:
                pass
            
            # Perform OCR if enabled
            if self.enable_ocr:
                try:
                    ocr_text = pytesseract.image_to_string(image)
                    content["text_content"] = ocr_text
                    content["image_metadata"]["ocr_text"] = ocr_text
                    content["image_metadata"]["ocr_performed"] = True
                except Exception as e:
                    content["image_metadata"]["ocr_error"] = str(e)
                    content["image_metadata"]["ocr_performed"] = False
            else:
                content["image_metadata"]["ocr_performed"] = False
            
            # Create HTML
            image_info = content["image_metadata"]
            html_content = f"""<div class='image-document'>
                <img src='{file_path}' alt='Extracted image' 
                     width='{image_info["width"]}' 
                     height='{image_info["height"]}' />
                <div class='image-metadata'>
                    <p><strong>Format:</strong> {image_info["format"]}</p>
                    <p><strong>Size:</strong> {image_info["width"]}x{image_info["height"]}</p>
                    <p><strong>Mode:</strong> {image_info["mode"]}</p>
                    {f'<p><strong>OCR Text:</strong> {content["text_content"]}</p>' if content["text_content"] else ''}
                </div>
            </div>"""
            
            content["html_content"] = html_content
            
            # Create markdown
            markdown_content = f"![Extracted image]({file_path})\n\n"
            if image_info:
                markdown_content += f"- **Format:** {image_info['format']}\n"
                markdown_content += f"- **Size:** {image_info['width']}x{image_info['height']}\n"
                markdown_content += f"- **Mode:** {image_info['mode']}\n"
            
            if content["text_content"]:
                markdown_content += f"\n- **OCR Text:** {content['text_content']}"
            
            content["markdown_content"] = markdown_content
            
            return content
        
        except Exception as e:
            logger.error(f"Failed to extract image content: {e}")
            return {
                "text_content": "",
                "html_content": "",
                "markdown_content": "",
                "extraction_method": "image",
                "success": False,
                "error_message": str(e)
            }
    
    # ==================== WEB CONTENT EXTRACTORS ====================
    
    def _extract_html(self, file_path: str) -> Dict[str, Any]:
        """Extract content from HTML file"""
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract title
            title = soup.title.string if soup.title else ""
            
            # Extract text content
            text_content = soup.get_text(separator=' ', strip=True)
            
            # Extract metadata
            metadata = {}
            
            # Meta tags
            for meta in soup.find_all('meta'):
                name = meta.get('name') or meta.get('property')
                content_attr = meta.get('content')
                if name and content_attr:
                    metadata[name] = content_attr
            
            # Clean HTML
            cleaned_html = str(soup)
            
            # Create markdown
            markdown_content = self._html_to_markdown(cleaned_html)
            
            return {
                "text_content": text_content,
                "html_content": cleaned_html,
                "markdown_content": markdown_content,
                "title": title,
                "document_metadata": metadata,
                "extraction_method": "html",
                "success": True
            }
        
        except Exception as e:
            logger.error(f"Failed to extract HTML content: {e}")
            return {
                "text_content": "",
                "html_content": "",
                "markdown_content": "",
                "extraction_method": "html",
                "success": False,
                "error_message": str(e)
            }
    
    def _extract_markdown(self, file_path: str) -> Dict[str, Any]:
        """Extract content from Markdown file"""
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                markdown_content = file.read()
            
            # Convert to HTML
            html_content = markdown.markdown(
                markdown_content,
                extensions=['codehilite', 'tables', 'toc']
            )
            
            # Extract title (first h1 or first line)
            title = ""
            lines = markdown_content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('# '):
                    title = line[2:].strip()
                    break
                elif line and not line.startswith('#') and not line.startswith('*') and not line.startswith('-'):
                    title = line
                    break
            
            # Extract text content (strip HTML)
            soup = BeautifulSoup(html_content, 'html.parser')
            text_content = soup.get_text(separator=' ', strip=True)
            
            return {
                "text_content": text_content,
                "html_content": html_content,
                "markdown_content": markdown_content,
                "title": title,
                "extraction_method": "markdown",
                "success": True
            }
        
        except Exception as e:
            logger.error(f"Failed to extract Markdown content: {e}")
            return {
                "text_content": "",
                "html_content": "",
                "markdown_content": "",
                "extraction_method": "markdown",
                "success": False,
                "error_message": str(e)
            }
    
    # ==================== GENERIC EXTRACTORS ====================
    
    def _extract_generic_text(self, file_path: str) -> Dict[str, Any]:
        """Generic text extraction using textract"""
        
        try:
            # Try textract for generic extraction
            text_content = textract.process(file_path).decode('utf-8')
            
            # Create HTML
            html_content = f"<div class='generic-document'>\n<pre>{text_content}</pre>\n</div>"
            
            # Create markdown
            markdown_content = text_content
            
            return {
                "text_content": text_content,
                "html_content": html_content,
                "markdown_content": markdown_content,
                "extraction_method": "generic",
                "success": True
            }
        
        except Exception as e:
            logger.error(f"Failed generic text extraction: {e}")
            return {
                "text_content": "",
                "html_content": "",
                "markdown_content": "",
                "extraction_method": "generic",
                "success": False,
                "error_message": str(e)
            }
    
    # ==================== UTILITY METHODS ====================
    
    def _html_to_markdown(self, html_content: str) -> str:
        """Convert HTML to Markdown"""
        
        try:
            # Try markdownify if available
            try:
                import markdownify
                return markdownify.markdownify(html_content)
            except ImportError:
                pass
            
            # Fallback to basic conversion
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Simple conversion rules
            for h1 in soup.find_all('h1'):
                h1.replace_with(f"# {h1.get_text()}")
            
            for h2 in soup.find_all('h2'):
                h2.replace_with(f"## {h2.get_text()}")
            
            for h3 in soup.find_all('h3'):
                h3.replace_with(f"### {h3.get_text()}")
            
            for p in soup.find_all('p'):
                p.replace_with(f"{p.get_text()}\n\n")
            
            for li in soup.find_all('li'):
                li.replace_with(f"- {li.get_text()}\n")
            
            for br in soup.find_all('br'):
                br.replace_with("\n")
            
            return soup.get_text()
        
        except Exception as e:
            logger.error(f"Failed to convert HTML to Markdown: {e}")
            return html_content
    
    async def batch_extract_content(self, 
                                   files: List[Dict[str, Any]]) -> List[ExtractedContent]:
        """Extract content from multiple files in batch"""
        
        try:
            # Process files in parallel
            tasks = []
            for file_info in files:
                task = self.extract_content(
                    file_id=file_info["file_id"],
                    file_path=file_info["file_path"],
                    file_name=file_info.get("file_name")
                )
                tasks.append(task)
            
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Process results
            extracted_contents = []
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    logger.error(f"Failed to extract content from file {i}: {result}")
                    extracted_contents.append(ExtractedContent(
                        file_id=files[i]["file_id"],
                        file_name=files[i].get("file_name", ""),
                        file_path=files[i]["file_path"],
                        mime_type="",
                        file_size=0,
                        text_content="",
                        html_content="",
                        markdown_content="",
                        success=False,
                        error_message=str(result)
                    ))
                else:
                    extracted_contents.append(result)
            
            return extracted_contents
        
        except Exception as e:
            logger.error(f"Failed batch content extraction: {e}")
            return []
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        
        return list(self.format_handlers.keys())
    
    def is_format_supported(self, file_ext: str) -> bool:
        """Check if file format is supported"""
        
        return file_ext.lower() in self.format_handlers

# Global content extractor instance
_content_extractor: Optional[ContentExtractor] = None

async def get_content_extractor() -> Optional[ContentExtractor]:
    """Get global content extractor instance"""
    
    global _content_extractor
    
    if _content_extractor is None:
        try:
            config = get_config_instance()
            _content_extractor = ContentExtractor(config)
            logger.info("Content Extractor created")
        except Exception as e:
            logger.error(f"Failed to create Content Extractor: {e}")
            _content_extractor = None
    
    return _content_extractor

def clear_content_extractor():
    """Clear global content extractor instance"""
    
    global _content_extractor
    _content_extractor = None
    logger.info("Content Extractor cleared")

# Export classes and functions
__all__ = [
    'ContentExtractor',
    'ExtractedContent',
    'get_content_extractor',
    'clear_content_extractor'
]